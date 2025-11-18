from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, validator
from typing import List, Optional, Dict, Any
import os
import pandas as pd
import json
import logging
from dotenv import load_dotenv
import asyncio
import concurrent.futures
from functools import partial
import sqlite3  # Add this import for database access
import urllib.parse
from datetime import datetime
from typing import Union
from starlette.exceptions import HTTPException as StarletteHTTPException
from docx import Document as DocxDocument

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(title="Financial Data API", version="1.0.0", docs_url="/api/docs", redoc_url="/api/redoc")

# Add CORS middleware with more permissive settings
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080", "http://localhost:8081", "http://127.0.0.1:8080", "http://127.0.0.1:8081", "http://localhost:5173", "https://localhost", "http://localhost:9000", "http://127.0.0.1:9000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    allow_origin_regex="https?://.*",
    expose_headers=["*"]
)



# Pydantic models for request/response
class ExcelDataResponse(BaseModel):
    data: List[Dict[str, Any]]
    columns: List[str]
    count: int



class SEBIExcelSummary(BaseModel):
    id: int
    date_key: str
    row_index: int
    pdf_link: Optional[str]
    summary: Optional[str]
    inserted_at: str
    entity_name: Optional[str] = None
    nature: Optional[str] = None

class SEBIAnalysisDataResponse(BaseModel):
    data: List[SEBIExcelSummary]
    count: int

# Add Pydantic model for BSE alerts data
class BSEAlertsDataResponse(BaseModel):
    data: List[Dict[str, Any]]
    count: int

# Add Pydantic models for admin authentication
class AdminLoginRequest(BaseModel):
    username: str
    password: str

class AdminLoginResponse(BaseModel):
    success: bool
    message: str
    token: Optional[str] = None

class AdminCredentialsResponse(BaseModel):
    id: int
    username: str



# Add Pydantic models for email management
class EmailEntry(BaseModel):
    email: str

class EmailResponse(BaseModel):
    email: str

class EmailListResponse(BaseModel):
    emails: List[str]
    count: int

# Add Pydantic models for visit tracking
class VisitCountResponse(BaseModel):
    count: int
    message: str

class VisitIncrementResponse(BaseModel):
    success: bool
    new_count: int
    message: str

# Add Pydantic models for Directors' Disclosure
class DirectorMasterResponse(BaseModel):
    id: int
    name: str
    din: str
    created_at: str

class DirectorsMasterResponse(BaseModel):
    data: List[DirectorMasterResponse]
    count: int

class DisclosureResponse(BaseModel):
    id: int
    director_name: str
    din: str
    disclosure_date: str
    disclosure_type: str
    file_path: str

class DisclosuresResponse(BaseModel):
    data: List[DisclosureResponse]
    count: int

class DisclosureContentResponse(BaseModel):
    content: str

class DisclosureAnalyticsResponse(BaseModel):
    total_disclosures: int
    by_type: List[Dict[str, Any]]
    by_month: List[Dict[str, Any]]
    by_director: List[Dict[str, Any]]





# Initialize visits database
def init_visits_db():
    """Initialize the visits database with a visits table"""
    db_path = os.path.join(os.path.dirname(__file__), "public", "visits.db")
    
    # Create database and table if they don't exist
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Create visits table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS visits (
            id INTEGER PRIMARY KEY,
            count INTEGER DEFAULT 0,
            last_updated TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    
    # Initialize with a default row if table is empty
    cursor.execute("SELECT COUNT(*) FROM visits")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO visits (count) VALUES (0)")
    
    conn.commit()
    conn.close()

# Call init_visits_db on startup
@app.on_event("startup")
async def startup_event():
    init_visits_db()

# Add endpoint to get visit count
@app.get("/visits/count", response_model=VisitCountResponse)
async def get_visit_count():
    """Get the current visit count"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "visits.db")
        
        def fetch_visit_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT count FROM visits WHERE id = 1")
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else 0
        
        loop = asyncio.get_event_loop()
        count = await loop.run_in_executor(thread_pool, fetch_visit_count)
        
        return VisitCountResponse(
            count=count,
            message="Successfully retrieved visit count"
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching visit count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch visit count: {error_message}")

# Add endpoint to increment visit count
@app.post("/visits/increment", response_model=VisitIncrementResponse)
async def increment_visit_count():
    """Increment the visit count by 1"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "visits.db")
        
        def update_visit_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("UPDATE visits SET count = count + 1, last_updated = CURRENT_TIMESTAMP WHERE id = 1")
            conn.commit()
            
            # Get the new count
            cursor.execute("SELECT count FROM visits WHERE id = 1")
            result = cursor.fetchone()
            new_count = result[0] if result else 0
            
            conn.close()
            return new_count
        
        loop = asyncio.get_event_loop()
        new_count = await loop.run_in_executor(thread_pool, update_visit_count)
        
        return VisitIncrementResponse(
            success=True,
            new_count=new_count,
            message="Successfully incremented visit count"
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error incrementing visit count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to increment visit count: {error_message}")

async def read_excel_sheet(file_path: str, sheet_name: str):
    """Read a specific sheet from an Excel file asynchronously using thread pool"""
    loop = asyncio.get_event_loop()
    func = partial(pd.read_excel, file_path, sheet_name=sheet_name)
    return await loop.run_in_executor(thread_pool, func)

# NOTE: The root endpoint (/) is intentionally not defined here to allow
# the React app to be served from the root path via static file serving.
# All API endpoints are available at their respective paths.

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "Financial Data API",
        "timestamp": logging.Formatter().formatTime(logging.LogRecord("", 0, "", 0, "", (), None))
    }

@app.get("/excel-data/{file_name}", response_model=ExcelDataResponse)
async def get_excel_data(file_name: str, sheet_name: str = "Sheet1"):
    """Get data from an Excel file"""
    try:
        # Define the path to the Excel file in the public folder
        excel_folder = os.path.join(os.path.dirname(__file__), "public", "excel")
        file_path = os.path.join(excel_folder, file_name)
        
        # Check if file exists in excel subdirectory
        if not os.path.exists(file_path):
            # Try with different case in excel subdirectory
            if file_name.lower() != file_name:
                file_path = os.path.join(excel_folder, file_name.lower())
            else:
                file_path = os.path.join(excel_folder, file_name.upper())
                
            # If still not found, check in the main public directory
            if not os.path.exists(file_path):
                # Check in main public directory
                public_folder = os.path.join(os.path.dirname(__file__), "public")
                file_path = os.path.join(public_folder, file_name)
                
                # Try with different case in main public directory
                if not os.path.exists(file_path):
                    if file_name.lower() != file_name:
                        file_path = os.path.join(public_folder, file_name.lower())
                    else:
                        file_path = os.path.join(public_folder, file_name.upper())
                
                if not os.path.exists(file_path):
                    raise HTTPException(status_code=404, detail=f"Excel file {file_name} not found")
        
        # Read the Excel file asynchronously
        df = await read_excel_sheet(file_path, sheet_name)
        
        # Convert to list of dictionaries
        data = df.to_dict('records')
        columns = df.columns.tolist()
        
        return ExcelDataResponse(
            data=data,
            columns=columns,
            count=len(data)
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error reading Excel file {file_name}: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to read Excel file: {error_message}")

# Add a new endpoint for BSE alerts data
@app.get("/bse-alerts", response_model=SEBIAnalysisDataResponse)
async def get_bse_alerts_data(limit: int = 100, offset: int = 0):
    """Get BSE alerts data from the notifications database"""
    try:
        # Define path to the notifications database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "notifications.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="BSE alerts database file not found")
        
        # Connect to the database and fetch data
        def fetch_bse_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # First, get the total count of records that match our criteria
            cursor.execute("""
                SELECT COUNT(*) 
                FROM DailyLogs 
                WHERE Link IS NOT NULL AND Link != 'NIL'
            """)
            total_count = cursor.fetchone()[0]
            
            # Fetch data from DailyLogs table with limit and offset for pagination
            # Only include records where Link is not NULL and not 'NIL'
            cursor.execute("""
                SELECT SrNo, EntityName, Link, Nature, Summary, Date 
                FROM DailyLogs 
                WHERE Link IS NOT NULL AND Link != 'NIL'
                ORDER BY Date DESC, SrNo ASC 
                LIMIT ? OFFSET ?
            """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Get column names
            column_names = [description[0] for description in cursor.description]
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                # Create a dictionary with the expected keys for the frontend
                record = dict(zip(column_names, row))
                
                # Rename keys to match the frontend expectations
                record['id'] = record.pop('SrNo', None)
                record['date_key'] = record.pop('Date', '')
                record['row_index'] = record.pop('SrNo', 0)  # Use SrNo as row_index
                record['pdf_link'] = record.pop('Link', '')
                record['summary'] = record.pop('Summary', '')
                record['inserted_at'] = record.pop('Date', '')  # Use Date as inserted_at
                # Preserve EntityName and Nature for BSE alerts
                if 'EntityName' in record:
                    record['entity_name'] = record.pop('EntityName')
                else:
                    record['entity_name'] = None
                if 'Nature' in record:
                    record['nature'] = record.pop('Nature')
                else:
                    record['nature'] = None
                
                data.append(record)
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_bse_data)
        
        return SEBIAnalysisDataResponse(
            data=data,
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching BSE alerts data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch BSE alerts data: {error_message}")

# Add endpoint for SEBI analysis data
@app.get("/sebi-analysis-data", response_model=SEBIAnalysisDataResponse)
async def get_sebi_excel_data(limit: int = 100, offset: int = 0):
    """Get SEBI analysis data from the SEBI database"""
    try:
        # Define the path to the SEBI database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "sebi_excel_master.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="SEBI database file not found")
        
        # Connect to the database and fetch data
        def fetch_sebi_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # First, get the total count of records
            cursor.execute("SELECT COUNT(*) FROM excel_summaries")
            total_count = cursor.fetchone()[0]
            
            # Fetch data from excel_summaries table with limit and offset for pagination
            cursor.execute("""
                SELECT id, date_key, row_index, pdf_link, summary, inserted_at 
                FROM excel_summaries 
                ORDER BY date_key DESC, row_index ASC 
                LIMIT ? OFFSET ?
            """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                record = {
                    'id': row[0],
                    'date_key': row[1],
                    'row_index': row[2],
                    'pdf_link': row[3],
                    'summary': row[4],
                    'inserted_at': row[5]
                }
                data.append(record)
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_sebi_data)
        
        return SEBIAnalysisDataResponse(
            data=data,
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching SEBI analysis data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch SEBI analysis data: {error_message}")

# Add endpoint for RBI analysis data
@app.get("/rbi-analysis-data", response_model=SEBIAnalysisDataResponse)
async def get_rbi_excel_data(limit: int = 100, offset: int = 0):
    """Get RBI analysis data from the RBI database"""
    try:
        # Define the path to the RBI database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "rbi.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="RBI database file not found")
        
        # Connect to the database and fetch data
        def fetch_rbi_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # First, get the total count of records that match our criteria
            cursor.execute("""
                SELECT COUNT(*) 
                FROM master_summaries 
                WHERE NOT (pdf_link = 'NIL' AND summary = 'NIL')
            """)
            total_count = cursor.fetchone()[0]
            
            # Fetch data from master_summaries table with limit and offset for pagination
            # Only exclude records where both pdf_link and summary are 'NIL'
            cursor.execute("""
                SELECT id, run_date, pdf_link, summary, created_at 
                FROM master_summaries 
                WHERE NOT (pdf_link = 'NIL' AND summary = 'NIL')
                ORDER BY run_date DESC, id ASC 
                LIMIT ? OFFSET ?
            """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                record = {
                    'id': row[0],
                    'date_key': row[1],  # Using date_key to match frontend expectations
                    'row_index': row[0],  # Using id as row_index
                    'pdf_link': row[2],
                    'summary': row[3],
                    'inserted_at': row[4]
                }
                data.append(record)
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_rbi_data)
        
        return SEBIAnalysisDataResponse(
            data=data,
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching RBI analysis data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch RBI analysis data: {error_message}")



# Add endpoint for BSE monthly analysis count
@app.get("/bse-monthly-count")
async def get_bse_monthly_count():
    """Get the count of BSE notifications for the current month"""
    try:
        # Define path to the notifications database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "notifications.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Notifications database file not found")
        
        # Connect to the database and fetch count
        def fetch_bse_monthly_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Get count of records for current month where Link is not NULL and not 'NIL'
            cursor.execute("""
                SELECT COUNT(*) 
                FROM DailyLogs 
                WHERE Date >= date('now', 'start of month') 
                AND Date < date('now', 'start of month', '+1 month')
                AND Link IS NOT NULL AND Link != 'NIL'
            """)
            
            count = cursor.fetchone()[0]
            conn.close()
            
            return {"count": count}
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(thread_pool, fetch_bse_monthly_count)
        
        return result
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching BSE monthly count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch BSE monthly count: {error_message}")

# Add a new endpoint for monthly BSE alerts count
@app.get("/api/bse-alerts-monthly-count")
async def get_bse_alerts_monthly_count():
    """Get monthly count of BSE alerts from the notifications database"""
    try:
        # Define path to the notifications database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "notifications.db")
       
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="BSE alerts database file not found")
       
        # Connect to the database and fetch data
        def fetch_counts():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
           
            # Get count of records grouped by month
            cursor.execute("""
                SELECT
                    strftime('%Y-%m', Date) as month,
                    COUNT(*) as count
                FROM DailyLogs
                WHERE Link IS NOT NULL AND Link != 'NIL'
                GROUP BY strftime('%Y-%m', Date)
                ORDER BY month DESC
            """)
           
            rows = cursor.fetchall()
           
            # Convert to list of dictionaries
            monthly_data = []
            for row in rows:
                monthly_data.append({
                    'month': row[0],
                    'count': row[1]
                })
           
            # Get total count of all BSE notifications
            cursor.execute("""
                SELECT COUNT(*)
                FROM DailyLogs
                WHERE Link IS NOT NULL AND Link != 'NIL'
            """)
           
            total_count = cursor.fetchone()[0]
           
            # Calculate average notifications per month
            average_count = 0
            if len(monthly_data) > 0:
                total_notifications = sum(item['count'] for item in monthly_data)
                average_count = round(total_notifications / len(monthly_data))
           
            conn.close()
            return monthly_data, total_count, average_count
       
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        monthly_data, total_count, average_count = await loop.run_in_executor(thread_pool, fetch_counts)
       
        return {"monthly_data": monthly_data, "total_count": total_count, "average_count": average_count}
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching BSE alerts monthly count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch BSE alerts monthly count: {error_message}")

@app.get("/api/bse-alerts-monthly-total")
async def get_bse_alerts_monthly_total():
    """Get total count of BSE alerts for the current month"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "notifications.db")

        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="BSE alerts database file not found")

        def fetch_total_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT COUNT(*)
                FROM DailyLogs
                WHERE Date >= date('now', 'start of month')
                AND Date < date('now', 'start of month', '+1 month')
                AND Link IS NOT NULL AND Link != 'NIL'
            """)

            count = cursor.fetchone()[0]
            conn.close()
            return count

        loop = asyncio.get_event_loop()
        total_count = await loop.run_in_executor(thread_pool, fetch_total_count)

        return {"count": total_count}
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching BSE alerts monthly total count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch BSE alerts monthly total count: {error_message}")

@app.get("/api/rbi-total-count")
async def get_rbi_total_count():
    """Get total count of RBI notifications"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "rbi.db")

        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="RBI database file not found")

        def fetch_total_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT COUNT(*) 
                FROM master_summaries 
                WHERE NOT (pdf_link = 'NIL' AND summary = 'NIL')
            """)

            count = cursor.fetchone()[0]
            conn.close()
            return count

        loop = asyncio.get_event_loop()
        total_count = await loop.run_in_executor(thread_pool, fetch_total_count)

        return {"count": total_count}
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching RBI total count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch RBI total count: {error_message}")

@app.get("/api/sebi-total-count")
async def get_sebi_total_count():
    """Get total count of SEBI notifications"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "sebi_excel_master.db")

        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="SEBI database file not found")

        def fetch_total_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()

            cursor.execute("SELECT COUNT(*) FROM excel_summaries")

            count = cursor.fetchone()[0]
            conn.close()
            return count

        loop = asyncio.get_event_loop()
        total_count = await loop.run_in_executor(thread_pool, fetch_total_count)

        return {"count": total_count}
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching SEBI total count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch SEBI total count: {error_message}")

# Add endpoint for SEBI analysis data
@app.get("/sebi-analysis-data", response_model=SEBIAnalysisDataResponse)
async def get_sebi_analysis_data(limit: int = 10, offset: int = 0):
    """Get SEBI analysis data with pagination"""
    try:
        # Define path to the database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "sebi_analysis.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="SEBI database file not found")
        
        # Connect to the database and fetch data
        def fetch_sebi_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # First, get the total count of records that match our criteria
            cursor.execute("""
                SELECT COUNT(*) 
                FROM master_summaries 
                WHERE NOT (pdf_link = 'NIL' AND summary = 'NIL')
            """)
            total_count = cursor.fetchone()[0]
            
            # Fetch data from master_summaries table with limit and offset for pagination
            # Only exclude records where both pdf_link and summary are 'NIL'
            cursor.execute("""
                SELECT id, run_date, pdf_link, summary, created_at 
                FROM master_summaries 
                WHERE NOT (pdf_link = 'NIL' AND summary = 'NIL')
                ORDER BY run_date DESC, id ASC 
                LIMIT ? OFFSET ?
            """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                record = {
                    'id': row[0],
                    'date_key': row[1],  # Using date_key to match frontend expectations
                    'row_index': row[0],  # Using id as row_index
                    'pdf_link': row[2],
                    'summary': row[3],
                    'inserted_at': row[4]
                }
                data.append(record)
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_sebi_data)
        
        return SEBIAnalysisDataResponse(
            data=data,
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching SEBI analysis data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch SEBI analysis data: {error_message}")

# Add admin authentication endpoints
@app.post("/admin/login", response_model=AdminLoginResponse)
async def admin_login(credentials: AdminLoginRequest):
    """Authenticate admin user"""
    try:
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and verify credentials
        def verify_admin_credentials():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Check if the provided credentials match any admin user
            cursor.execute("""
                SELECT id, username FROM admin_credentials 
                WHERE username = ? AND password = ?
            """, (credentials.username, credentials.password))
            
            result = cursor.fetchone()
            conn.close()
            
            return result
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        admin_user = await loop.run_in_executor(thread_pool, verify_admin_credentials)
        
        if admin_user:
            # In a real application, you would generate a proper JWT token
            # For now, we'll just return a success response
            return AdminLoginResponse(
                success=True,
                message="Login successful",
                token=f"admin_token_{admin_user[0]}"  # Simple token for demonstration
            )
        else:
            return AdminLoginResponse(
                success=False,
                message="Invalid credentials"
            )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error during admin login: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to authenticate: {error_message}")

# Add email management endpoints (admin only)
@app.get("/emails", response_model=EmailListResponse)
async def get_emails(search: Optional[str] = None):
    """Get all email addresses with optional search filter"""
    try:
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and fetch emails
        def fetch_emails():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Build query with optional search filter
            if search:
                # Case-insensitive search in email column
                query = "SELECT email FROM email WHERE LOWER(email) LIKE LOWER(?) ORDER BY email"
                search_pattern = f"%{search}%"
                cursor.execute(query, (search_pattern,))
            else:
                # Get all emails
                cursor.execute("SELECT email FROM email ORDER BY email")
            
            rows = cursor.fetchall()
            conn.close()
            
            return [row[0] for row in rows]
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        emails = await loop.run_in_executor(thread_pool, fetch_emails)
        
        return EmailListResponse(
            emails=emails,
            count=len(emails)
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching emails: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch emails: {error_message}")

@app.post("/emails", response_model=EmailResponse)
async def add_email(email_entry: EmailEntry):
    """Add a new email address (admin only)"""
    try:
        # Validate email format
        import re
        email_regex = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
        if not re.match(email_regex, email_entry.email):
            raise HTTPException(status_code=400, detail="Invalid email format")
        
        # Validate that email is from adani.com or pspprojects.com domain (case-insensitive)
        email_lower = email_entry.email.lower()
        if not email_lower.endswith('@adani.com') and not email_lower.endswith('@pspprojects.com'):
            raise HTTPException(status_code=400, detail="Only emails from adani.com or pspprojects.com domains are allowed")
        
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and add email
        def add_email_to_db():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Check if email already exists (case-insensitive)
            cursor.execute("SELECT COUNT(*) FROM email WHERE LOWER(email) = LOWER(?)", (email_entry.email,))
            count = cursor.fetchone()[0]
            
            if count > 0:
                conn.close()
                raise HTTPException(status_code=409, detail="Email already exists")
            
            # Insert new email
            cursor.execute("INSERT INTO email (email) VALUES (?)", (email_entry.email,))
            conn.commit()
            conn.close()
            
            return email_entry.email
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        email = await loop.run_in_executor(thread_pool, add_email_to_db)
        
        return EmailResponse(email=email)
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error adding email: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to add email: {error_message}")

@app.delete("/emails/{email_address}", response_model=EmailResponse)
async def delete_email(email_address: str):
    """Delete an email address (admin only)"""
    try:
        # Decode URL encoded email address
        email = urllib.parse.unquote(email_address)
        
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and delete email
        def delete_email_from_db():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Check if email exists
            cursor.execute("SELECT COUNT(*) FROM email WHERE email = ?", (email,))
            count = cursor.fetchone()[0]
            
            if count == 0:
                conn.close()
                raise HTTPException(status_code=404, detail="Email not found")
            
            # Delete the email
            cursor.execute("DELETE FROM email WHERE email = ?", (email,))
            conn.commit()
            conn.close()
            
            return email
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        deleted_email = await loop.run_in_executor(thread_pool, delete_email_from_db)
        
        return EmailResponse(email=deleted_email)
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error deleting email: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to delete email: {error_message}")

# Directors' Disclosure Endpoints
@app.get("/api/directors-master", response_model=DirectorsMasterResponse)
async def get_directors_master():
    """Get all directors from directors database"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "directors.db")
        
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Directors database not found")
        
        def fetch_directors():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT id, name, din, created_at FROM directors ORDER BY name")
            rows = cursor.fetchall()
            conn.close()
            
            return [{
                'id': row[0],
                'name': row[1],
                'din': row[2],
                'created_at': row[3]
            } for row in rows]
        
        loop = asyncio.get_event_loop()
        directors = await loop.run_in_executor(thread_pool, fetch_directors)
        
        return DirectorsMasterResponse(
            data=[DirectorMasterResponse(**d) for d in directors],
            count=len(directors)
        )
    except Exception as e:
        logger.error(f"Error fetching directors master: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch directors: {str(e)}")

@app.get("/api/directors-disclosures", response_model=DisclosuresResponse)
async def get_directors_disclosures():
    """Get all directors' disclosures from Word files"""
    try:
        # Path to disclosure output folder
        disclosures_dir = os.path.join(os.path.dirname(__file__), "public", "Directors Discloser Output")
        
        def fetch_disclosures():
            disclosures = []
            
            # Check if directory exists
            if not os.path.exists(disclosures_dir):
                logger.warning(f"Disclosures directory not found: {disclosures_dir}")
                return []
            
            # Scan directory for .docx files
            for idx, filename in enumerate(sorted(os.listdir(disclosures_dir))):
                if filename.endswith('.docx') and not filename.startswith('~$'):
                    file_path = os.path.join(disclosures_dir, filename)
                    
                    # Extract metadata from filename or file stats
                    file_stat = os.stat(file_path)
                    created_date = datetime.fromtimestamp(file_stat.st_mtime).strftime('%Y-%m-%d')
                    
                    # Extract director name from filename (remove _MBP.docx)
                    director_name = filename.replace('_MBP.docx', '').replace('.docx', '').strip()
                    din = 'N/A'  # Default value
                    
                    # Try to extract DIN from document
                    try:
                        doc = DocxDocument(file_path)
                        # Look for DIN in document paragraphs
                        for para in doc.paragraphs:
                            text = para.text.strip()
                            # Match pattern like "DIN : 12345678" or "DIN: 12345678"
                            import re
                            din_match = re.search(r'DIN\s*:\s*([0-9]{8})', text, re.IGNORECASE)
                            if din_match:
                                din = din_match.group(1)
                                break
                    except Exception as e:
                        logger.warning(f"Error reading DIN from {filename}: {e}")
                    
                    disclosures.append({
                        'id': idx + 1,
                        'director_name': director_name,
                        'din': din,
                        'disclosure_date': created_date,
                        'disclosure_type': 'MBP-1',
                        'file_path': filename
                    })
            
            return disclosures
        
        loop = asyncio.get_event_loop()
        disclosures = await loop.run_in_executor(thread_pool, fetch_disclosures)
        
        return DisclosuresResponse(
            data=[DisclosureResponse(**d) for d in disclosures],
            count=len(disclosures)
        )
    except Exception as e:
        logger.error(f"Error fetching disclosures: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch disclosures: {str(e)}")

@app.get("/api/directors-disclosures/{disclosure_id}/content", response_model=DisclosureContentResponse)
async def get_disclosure_content(disclosure_id: int):
    """Get content of a specific disclosure document"""
    try:
        disclosures_dir = os.path.join(os.path.dirname(__file__), "public", "Directors Discloser Output")
        
        def read_disclosure_content():
            if not os.path.exists(disclosures_dir):
                raise HTTPException(status_code=404, detail="Disclosures directory not found")
            
            # Get list of docx files
            docx_files = [f for f in os.listdir(disclosures_dir) 
                         if f.endswith('.docx') and not f.startswith('~$')]
            
            # Check if disclosure_id is valid
            if disclosure_id < 1 or disclosure_id > len(docx_files):
                raise HTTPException(status_code=404, detail="Disclosure not found")
            
            # Get the file at the specified index
            filename = sorted(docx_files)[disclosure_id - 1]
            file_path = os.path.join(disclosures_dir, filename)
            
            # Read Word document content
            try:
                doc = DocxDocument(file_path)
                
                # Extract all text from document
                content_parts = []
                
                # Add document title if available
                content_parts.append(f"Document: {filename}\n")
                content_parts.append("=" * 80 + "\n\n")
                
                # Extract all paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text + "\n")
                
                # Extract tables if any
                if doc.tables:
                    content_parts.append("\n" + "=" * 80 + "\n")
                    content_parts.append("TABLES\n")
                    content_parts.append("=" * 80 + "\n\n")
                    
                    for idx, table in enumerate(doc.tables):
                        content_parts.append(f"Table {idx + 1}:\n")
                        for row in table.rows:
                            row_text = " | ".join([cell.text.strip() for cell in row.cells])
                            content_parts.append(row_text + "\n")
                        content_parts.append("\n")
                
                full_content = "".join(content_parts)
                
                if not full_content.strip():
                    return "No content found in document"
                
                return full_content
                
            except Exception as e:
                logger.error(f"Error reading Word document: {e}")
                raise HTTPException(status_code=500, detail=f"Error reading document: {str(e)}")
        
        loop = asyncio.get_event_loop()
        content = await loop.run_in_executor(thread_pool, read_disclosure_content)
        
        return DisclosureContentResponse(content=content)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching disclosure content: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch content: {str(e)}")

@app.get("/api/directors-disclosures/analytics", response_model=DisclosureAnalyticsResponse)
async def get_disclosures_analytics():
    """Get analytics data for directors' disclosures"""
    try:
        disclosures_dir = os.path.join(os.path.dirname(__file__), "public", "Directors Discloser Output")
        
        def calculate_analytics():
            from collections import defaultdict
            import re
            
            if not os.path.exists(disclosures_dir):
                # Return empty analytics if directory doesn't exist
                return {
                    'total_disclosures': 0,
                    'by_type': [],
                    'by_month': [],
                    'by_director': []
                }
            
            docx_files = [f for f in os.listdir(disclosures_dir) 
                         if f.endswith('.docx') and not f.startswith('~$')]
            
            total_count = len(docx_files)
            
            # Track statistics
            by_type = defaultdict(int)
            by_month = defaultdict(int)
            by_director = defaultdict(int)
            
            for filename in docx_files:
                file_path = os.path.join(disclosures_dir, filename)
                
                # Get file modification date for monthly stats
                file_stat = os.stat(file_path)
                file_date = datetime.fromtimestamp(file_stat.st_mtime)
                month_key = file_date.strftime('%b %Y')
                by_month[month_key] += 1
                
                # Extract director name from filename (remove _MBP.docx)
                director_name = filename.replace('_MBP.docx', '').replace('.docx', '').strip()
                
                # Try to read document for better classification
                try:
                    doc = DocxDocument(file_path)
                    
                    # Look for disclosure type in content
                    for para in doc.paragraphs[:15]:
                        text = para.text.lower()
                        
                        # Classify disclosure type
                        if 'shareholding' in text or 'shares' in text:
                            by_type['Shareholding'] += 1
                            break
                        elif 'transaction' in text or 'acquisition' in text:
                            by_type['Transaction'] += 1
                            break
                        elif 'interest' in text or 'concern' in text:
                            by_type['Interest'] += 1
                            break
                    else:
                        # Default type - MBP-1 form
                        by_type['MBP-1'] += 1
                    
                except Exception as e:
                    logger.warning(f"Error analyzing {filename}: {e}")
                    by_type['MBP-1'] += 1
                
                # Track by director
                by_director[director_name] += 1
            
            # Convert to list format for response
            analytics = {
                'total_disclosures': total_count,
                'by_type': [{'type': k, 'count': v} for k, v in sorted(by_type.items(), key=lambda x: -x[1])],
                'by_month': [{'month': k, 'count': v} for k, v in sorted(by_month.items())],
                'by_director': [{'director': k, 'count': v} for k, v in sorted(by_director.items(), key=lambda x: -x[1])[:10]]  # Top 10
            }
            
            return analytics
        
        loop = asyncio.get_event_loop()
        analytics = await loop.run_in_executor(thread_pool, calculate_analytics)
        
        return DisclosureAnalyticsResponse(**analytics)
    except Exception as e:
        logger.error(f"Error fetching analytics: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch analytics: {str(e)}")

# Minutes Preparation - Directors Endpoint
@app.get("/directors", response_model=DirectorsMasterResponse)
async def get_directors_for_minutes():
    """Get all directors from directors database for Minutes Preparation"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "directors.db")
        
        if not os.path.exists(db_path):
            logger.warning(f"Directors database not found: {db_path}")
            return DirectorsMasterResponse(data=[], count=0)
        
        def fetch_directors():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT id, name, din, created_at FROM directors ORDER BY name")
            rows = cursor.fetchall()
            conn.close()
            
            return [{
                'id': row[0],
                'name': row[1],
                'din': row[2],
                'created_at': row[3]
            } for row in rows]
        
        loop = asyncio.get_event_loop()
        directors = await loop.run_in_executor(thread_pool, fetch_directors)
        
        return DirectorsMasterResponse(
            data=[DirectorMasterResponse(**d) for d in directors],
            count=len(directors)
        )
    except Exception as e:
        logger.error(f"Error fetching directors for minutes: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch directors: {str(e)}")

# Places Management - for Meeting Places
class PlaceResponse(BaseModel):
    id: int
    name: str
    address: str
    is_default: bool
    created_at: str

class PlacesListResponse(BaseModel):
    data: List[PlaceResponse]
    count: int

class PlaceCreateRequest(BaseModel):
    name: str
    address: str
    is_default: bool = False

def init_places_db():
    """Initialize places database with default Adani Corporate House"""
    db_path = os.path.join(os.path.dirname(__file__), "public", "places.db")
    
    # Create public directory if it doesn't exist
    os.makedirs(os.path.dirname(db_path), exist_ok=True)
    
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Create places table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS places (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            address TEXT NOT NULL,
            is_default BOOLEAN DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Check if default place exists
    cursor.execute("SELECT COUNT(*) FROM places WHERE is_default = 1")
    if cursor.fetchone()[0] == 0:
        # Insert default Adani Corporate House
        cursor.execute('''
            INSERT INTO places (name, address, is_default)
            VALUES (?, ?, ?)
        ''', (
            'Adani Corporate House',
            'Adani Corporate House, Shantigram, Near Vaishno Devi Circle, S. G. Highway, Khodiyar, Ahmedabad - 382421, Gujarat, India',
            1
        ))
    
    conn.commit()
    conn.close()

# Initialize places database on startup
init_places_db()

@app.get("/places", response_model=PlacesListResponse)
async def get_places():
    """Get all places from database"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "places.db")
        
        def fetch_places():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT id, name, address, is_default, created_at FROM places ORDER BY is_default DESC, name")
            rows = cursor.fetchall()
            conn.close()
            
            places = [
                PlaceResponse(
                    id=row[0],
                    name=row[1],
                    address=row[2],
                    is_default=bool(row[3]),
                    created_at=row[4]
                )
                for row in rows
            ]
            return places
        
        loop = asyncio.get_event_loop()
        places = await loop.run_in_executor(thread_pool, fetch_places)
        
        return PlacesListResponse(
            data=places,
            count=len(places)
        )
    except Exception as e:
        logger.error(f"Error fetching places: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch places: {str(e)}")

@app.post("/places", response_model=PlaceResponse)
async def create_place(request: PlaceCreateRequest):
    """Create a new place"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "places.db")
        
        def insert_place():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # If this is set as default, unset other defaults
            if request.is_default:
                cursor.execute("UPDATE places SET is_default = 0")
            
            cursor.execute('''
                INSERT INTO places (name, address, is_default)
                VALUES (?, ?, ?)
            ''', (request.name, request.address, request.is_default))
            
            place_id = cursor.lastrowid
            conn.commit()
            
            # Fetch the created place
            cursor.execute("SELECT id, name, address, is_default, created_at FROM places WHERE id = ?", (place_id,))
            row = cursor.fetchone()
            conn.close()
            
            return PlaceResponse(
                id=row[0],
                name=row[1],
                address=row[2],
                is_default=bool(row[3]),
                created_at=row[4]
            )
        
        loop = asyncio.get_event_loop()
        new_place = await loop.run_in_executor(thread_pool, insert_place)
        
        return new_place
    except Exception as e:
        logger.error(f"Error creating place: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to create place: {str(e)}")

# Minutes Generation Endpoint
class MinutesGenerationRequest(BaseModel):
    template: str
    companyName: str
    meetingNumber: str
    meetingType: str
    meetingDay: str
    meetingDate: str
    meetingStartTime: str
    meetingEndTime: str
    meetingPlace: str
    chairmanName: str
    presentDirectors: List[Dict[str, str]]
    inAttendance: List[Dict[str, str]]
    companySecretary: str
    previousMeetingDate: str
    authorisedOfficer: str
    quorum: str
    concerns: str
    declarations: str
    auditorPaymentAmount: str
    auditorPaymentWords: str
    financialYear: str
    agmNumber: str
    agmDay: str
    agmMonthName: str
    agmDate: str
    agmTime: str
    agmPlace: str
    recordingDate: str
    signingDate: str
    signingPlace: str

@app.post("/generate-minutes")
async def generate_minutes(request: MinutesGenerationRequest):
    """Generate meeting minutes document from template"""
    try:
        logger.info(f"Generating minutes for template: {request.template}")
        
        # Define template path
        template_path = os.path.join(os.path.dirname(__file__), "public", "templates", f"{request.template.lower()}_meeting_template.docx")
        
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail=f"Template {request.template} not found")
        
        def generate_document():
            # Load the template
            doc = DocxDocument(template_path)
            
            # Create placeholder mapping
            placeholders = {
                '[No. of Meeting]': request.meetingNumber,
                '[Type of Meeting]': request.meetingType,
                '[Name of Company]': request.companyName,
                '[Day of Meeting]': request.meetingDay,
                '[Date of Meeting]': request.meetingDate,
                '[Time: COMMENCED AT]': request.meetingStartTime,
                '[Time: CONCLUDED AT]': request.meetingEndTime,
                '[Place of Meeting]': request.meetingPlace,
                '[Chairman]': request.chairmanName,
                '[Director]': request.chairmanName,  # Same as chairman for now
                '[Date-previous-meeting]': request.previousMeetingDate,
                '[amount]': request.auditorPaymentAmount,
                '[Amount-in-words]': request.auditorPaymentWords,
                '[amount-in-words]': request.auditorPaymentWords,
                '[Year]': request.financialYear,
                '[year]': request.financialYear,
                '[YEAR]': request.financialYear,
                '[start-year]': request.financialYear.split('-')[0] if '-' in request.financialYear else request.financialYear,
                '[end-year]': request.financialYear.split('-')[1] if '-' in request.financialYear else str(int(request.financialYear) + 1),
                '[Day-of-meeting]': request.meetingDay,
                '[Month-of-meeting]': request.agmMonthName,
                '[TIME]': request.agmTime,
                '[Office-address]': request.agmPlace,
                '[Date-of-Recording]': request.recordingDate,
                '[Date-of-signing]': request.signingDate,
                '[Place of signing]': request.signingPlace,
                '[Officer]': request.companySecretary or request.authorisedOfficer,
            }
            
            # Add directors (multiple instances)
            if request.presentDirectors and len(request.presentDirectors) > 0:
                for i, director in enumerate(request.presentDirectors[:4]):  # Up to 4 directors
                    placeholders[f'[Dir-name]'] = director.get('name', '')
                    placeholders[f'[Din-num]'] = director.get('din', '')
            
            # Replace placeholders in paragraphs
            for para in doc.paragraphs:
                for placeholder, value in placeholders.items():
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(value))
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, value in placeholders.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"meeting_minutes_{request.template}_{timestamp}.docx"
            output_path = os.path.join(os.path.dirname(__file__), "public", "templates", filename)
            
            # Save the document
            doc.save(output_path)
            
            return filename, output_path
        
        # Run document generation in thread pool
        loop = asyncio.get_event_loop()
        filename, output_path = await loop.run_in_executor(thread_pool, generate_document)
        
        # Return file for download
        return FileResponse(
            path=output_path,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating minutes: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate minutes: {str(e)}")





# Custom thread pool for handling blocking operations
thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

# Custom static files class to handle SPA routing
class SPAStaticFiles(StaticFiles):
    async def get_response(self, path: str, scope):
        # Normalize path to check for API routes
        normalized_path = path.lstrip("/")
        
        # Skip API routes entirely - let FastAPI handle them
        # FastAPI routes are checked before mounted static files, but this is a safety check
        if normalized_path.startswith("api/"):
            raise StarletteHTTPException(status_code=404)
        
        try:
            return await super().get_response(path, scope)
        except (StarletteHTTPException, HTTPException) as ex:
            if ex.status_code == 404:
                # Skip API routes - don't serve index.html for them
                normalized_path = path.lstrip("/")
                if normalized_path.startswith("api/"):
                    raise ex
                # Return index.html for any non-existent file (SPA routing)
                return await super().get_response("index.html", scope)
            else:
                raise ex

# Serve static files from the dist directory (React build) - this must be the last route
# Only add this if the dist directory exists
DIST_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "dist")
if os.path.exists(DIST_DIR):
    
    app.mount("/", SPAStaticFiles(directory=DIST_DIR, html=True), name="static")

# Add a test endpoint to verify static file serving is working
@app.get("/test-static")
async def test_static_serving():
    """Test endpoint to verify static file serving is working"""
    return {"message": "Static file serving should be working correctly"}

if __name__ == "__main__":
    import uvicorn
    
    # Run without SSL on port 8001
    uvicorn.run(app, host="0.0.0.0", port=8000)