from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, validator
from typing import List, Optional, Dict, Any
import os
import pandas as pd
import json
import logging
from dotenv import load_dotenv
import asyncio
import concurrent.futures
from functools import partial
import sqlite3  # Add this import for database access
import urllib.parse
from docx import Document
import re
from datetime import datetime
from typing import Union

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(title="Financial Data API", version="1.0.0", docs_url="/api/docs", redoc_url="/api/redoc")

# Add CORS middleware with more permissive settings
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080", "http://localhost:8081", "http://127.0.0.1:8080", "http://127.0.0.1:8081", "http://localhost:5173", "https://localhost", "http://localhost:9000", "http://127.0.0.1:9000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    allow_origin_regex="https?://.*",
    expose_headers=["*"]
)

# Serve static files from the public/templates directory
templates_dir = os.path.join(os.path.dirname(__file__), "public", "templates")
if os.path.exists(templates_dir):
    app.mount("/templates", StaticFiles(directory=templates_dir), name="templates")

# Pydantic models for request/response
class ExcelDataResponse(BaseModel):
    data: List[Dict[str, Any]]
    columns: List[str]
    count: int

# Add new Pydantic model for directors
class DirectorResponse(BaseModel):
    id: int
    name: str
    din: str
    created_at: str

class DirectorsResponse(BaseModel):
    data: List[DirectorResponse]
    count: int

class SEBIExcelSummary(BaseModel):
    id: int
    date_key: str
    row_index: int
    pdf_link: Optional[str]
    summary: Optional[str]
    inserted_at: str
    entity_name: Optional[str] = None
    nature: Optional[str] = None

class SEBIAnalysisDataResponse(BaseModel):
    data: List[SEBIExcelSummary]
    count: int

# Add Pydantic models for admin authentication
class AdminLoginRequest(BaseModel):
    username: str
    password: str

class AdminLoginResponse(BaseModel):
    success: bool
    message: str
    token: Optional[str] = None

class AdminCredentialsResponse(BaseModel):
    id: int
    username: str

# Add Pydantic models for place management
class PlaceResponse(BaseModel):
    id: int
    name: str
    address: str
    is_default: bool
    created_at: str

class PlacesResponse(BaseModel):
    data: List[PlaceResponse]
    count: int

class PlaceCreateRequest(BaseModel):
    name: str
    address: str
    is_default: bool = False

class PlaceUpdateRequest(BaseModel):
    name: Optional[str] = None
    address: Optional[str] = None
    is_default: Optional[bool] = None

    created_at: str

# Add Pydantic models for email management
class EmailEntry(BaseModel):
    email: str

class EmailResponse(BaseModel):
    email: str

class EmailListResponse(BaseModel):
    emails: List[str]
    count: int

# Add Pydantic models for visit tracking
class VisitCountResponse(BaseModel):
    count: int
    message: str

class VisitIncrementResponse(BaseModel):
    success: bool
    new_count: int
    message: str

# Add Pydantic models for minutes generation
class Director(BaseModel):
    name: str
    din: str

    class Config:
        # Ensure validation runs even when values come from dict
        validate_assignment = True

class Attendee(BaseModel):
    name: str
    role: str

class Signatory(BaseModel):
    name: str
    role: str
    din: str

class YearRange(BaseModel):
    from_year: int
    to_year: int

class Q1MinutesData(BaseModel):
    # Company & meeting header
    companyName: str
    meetingNumber: str
    meetingType: str
    meetingDay: str
    meetingDate: str
    timeCommenced: str
    timeConcluded: str
    meetingPlace: str
    
    # Attendance block
    presentDirectors: List[Director]
    chairmanName: str
    inAttendance: List[Attendee]
    
    # Quorum & minutes confirmation
    previousMinutesDate: str
    
    # Disclosures under the Companies Act
    interestDisclosures: List[Director]
    disqualificationDeclarations: List[Director]
    
    # Statutory auditor's payment
    auditorPaymentNumber: int
    auditorPaymentWords: str
    auditorPaymentYear: int
    
    # Financial statements approval
    fsYear: int
    rptFinYearRange: YearRange
    signatory1: Signatory
    signatory2: Signatory
    
    # Directors' Report approval
    directorsReportYear: int
    
    # AGM notice & meeting details
    agmNumber: str
    agmDayName: str
    agmDay: int
    agmMonth: str
    agmYear: int
    agmTime: str
    registeredOfficeAddress: str
    chairmanShortName: str
    
    # Sign-off block
    recordingDate: str
    signingDate: str
    signingPlace: str
    signingChairmanName: str
    
    # Template selection
    template: str = "Q1"
    
    class Config:
        # Ensure validation runs even when values come from dict
        validate_assignment = True

class MinutesData(BaseModel):
    template: str  # New field for template selection
    companyName: str
    meetingNumber: str
    meetingType: str
    meetingDay: str
    meetingDate: str
    meetingStartTime: str
    meetingEndTime: str
    meetingPlace: str
    chairmanName: str = "Chairman Name"  # Provide default value
    directors: List[Director] = []  # Make optional with default
    inAttendance: List[Attendee] = []  # Add inAttendance field
    presentDirectors: List[Director] = []  # Add presentDirectors field
    authorisedOfficer: str = "Authorised Officer"  # Provide default value
    previousMeetingDate: str
    auditorPaymentAmount: str = "0"  # Provide default value
    auditorPaymentWords: str = "Amount in words"  # Provide default value
    financialYear: str = str(datetime.now().year)  # Provide default value
    agmNumber: str = ""  # Make optional
    agmDay: str = "1"  # Provide default value
    agmDate: str = datetime.now().strftime("%Y-%m-%d")  # Provide default value
    agmTime: str = ""  # Make optional
    agmPlace: str = ""  # Make optional
    recordingDate: str = ""  # Make optional
    signingDate: str = ""  # Make optional
    signingPlace: str = ""  # Add signingPlace field
    quorum: str = "Quorum details"  # Provide default value
    previousMinutes: str = "Previous minutes details"  # Provide default value
    concerns: str = "Concerns details"  # Provide default value
    declarations: str = "Declarations details"  # Provide default value
    auditorPayment: str = "Auditor payment details"  # Provide default value
    financialStatements: str = "Financial statements details"  # Provide default value
    directorsReport: str = "Directors report details"
    
    class Config:
        # Ensure validation runs even when values come from dict
        validate_assignment = True
    
    # Add validation for template
    @validator('template')
    def validate_template(cls, v):
        allowed_templates = ['Q1', 'Q2', 'Q3', 'Q4']
        if v not in allowed_templates:
            raise ValueError(f'Template must be one of: {", ".join(allowed_templates)}')
        return v
    
    # Add validation
    @validator('financialYear')
    def validate_financial_year(cls, v):
        if v and not re.match(r'^\d{4}$', v):
            raise ValueError('Financial year must be a 4-digit number')
        return v or str(datetime.now().year)  # Provide default if empty
    
    @validator('meetingDate', 'previousMeetingDate', 'agmDate', 'recordingDate', 'signingDate')
    def validate_dates(cls, v):
        if v:
            try:
                datetime.strptime(v, '%Y-%m-%d')
            except ValueError:
                raise ValueError('Date must be in YYYY-MM-DD format')
        return v or datetime.now().strftime("%Y-%m-%d")  # Provide default if empty
    
    @validator('directors', 'presentDirectors')
    def validate_directors(cls, v):
        # Allow empty directors list and provide default if needed
        if not v:
            return [Director(name="Director Name", din="12345678")]
        # Filter out empty directors
        valid_directors = [d for d in v if d.name.strip() and d.din.strip()]
        if not valid_directors:
            return [Director(name="Director Name", din="12345678")]
        return valid_directors

    @validator('companyName')
    def validate_company_name(cls, v):
        if not v or not v.strip():
            raise ValueError('Company name cannot be empty')
        return v.strip()

    @validator('chairmanName')
    def validate_chairman_name(cls, v):
        return v or "Chairman Name"  # Provide default if empty

    @validator('meetingNumber')
    def validate_meeting_number(cls, v):
        if not v or not v.strip():
            raise ValueError('Meeting number cannot be empty')
        return v.strip()

# Initialize visits database
def init_visits_db():
    """Initialize the visits database with a visits table"""
    db_path = os.path.join(os.path.dirname(__file__), "public", "visits.db")
    
    # Create database and table if they don't exist
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Create visits table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS visits (
            id INTEGER PRIMARY KEY,
            count INTEGER DEFAULT 0,
            last_updated TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    
    # Initialize with a default row if table is empty
    cursor.execute("SELECT COUNT(*) FROM visits")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO visits (count) VALUES (0)")
    
    conn.commit()
    conn.close()

# Call init_visits_db on startup
@app.on_event("startup")
async def startup_event():
    init_visits_db()

# Add endpoint to get visit count
@app.get("/visits/count", response_model=VisitCountResponse)
async def get_visit_count():
    """Get the current visit count"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "visits.db")
        
        def fetch_visit_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT count FROM visits WHERE id = 1")
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else 0
        
        loop = asyncio.get_event_loop()
        count = await loop.run_in_executor(thread_pool, fetch_visit_count)
        
        return VisitCountResponse(
            count=count,
            message="Successfully retrieved visit count"
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching visit count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch visit count: {error_message}")

# Add endpoint to increment visit count
@app.post("/visits/increment", response_model=VisitIncrementResponse)
async def increment_visit_count():
    """Increment the visit count by 1"""
    try:
        db_path = os.path.join(os.path.dirname(__file__), "public", "visits.db")
        
        def update_visit_count():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute("UPDATE visits SET count = count + 1, last_updated = CURRENT_TIMESTAMP WHERE id = 1")
            conn.commit()
            
            # Get the new count
            cursor.execute("SELECT count FROM visits WHERE id = 1")
            result = cursor.fetchone()
            new_count = result[0] if result else 0
            
            conn.close()
            return new_count
        
        loop = asyncio.get_event_loop()
        new_count = await loop.run_in_executor(thread_pool, update_visit_count)
        
        return VisitIncrementResponse(
            success=True,
            new_count=new_count,
            message="Successfully incremented visit count"
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error incrementing visit count: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to increment visit count: {error_message}")

async def read_excel_sheet(file_path: str, sheet_name: str):
    """Read a specific sheet from an Excel file asynchronously using thread pool"""
    loop = asyncio.get_event_loop()
    func = partial(pd.read_excel, file_path, sheet_name=sheet_name)
    return await loop.run_in_executor(thread_pool, func)

# NOTE: The root endpoint (/) is intentionally not defined here to allow
# the React app to be served from the root path via static file serving.
# All API endpoints are available at their respective paths.

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "Financial Data API",
        "timestamp": logging.Formatter().formatTime(logging.LogRecord("", 0, "", 0, "", (), None))
    }

@app.get("/excel-data/{file_name}", response_model=ExcelDataResponse)
async def get_excel_data(file_name: str, sheet_name: str = "Sheet1"):
    """Get data from an Excel file"""
    try:
        # Define the path to the Excel file in the public folder
        excel_folder = os.path.join(os.path.dirname(__file__), "public", "excel")
        file_path = os.path.join(excel_folder, file_name)
        
        # Check if file exists in excel subdirectory
        if not os.path.exists(file_path):
            # Try with different case in excel subdirectory
            if file_name.lower() != file_name:
                file_path = os.path.join(excel_folder, file_name.lower())
            else:
                file_path = os.path.join(excel_folder, file_name.upper())
                
            # If still not found, check in the main public directory
            if not os.path.exists(file_path):
                # Check in main public directory
                public_folder = os.path.join(os.path.dirname(__file__), "public")
                file_path = os.path.join(public_folder, file_name)
                
                # Try with different case in main public directory
                if not os.path.exists(file_path):
                    if file_name.lower() != file_name:
                        file_path = os.path.join(public_folder, file_name.lower())
                    else:
                        file_path = os.path.join(public_folder, file_name.upper())
                
                if not os.path.exists(file_path):
                    raise HTTPException(status_code=404, detail=f"Excel file {file_name} not found")
        
        # Read the Excel file asynchronously
        df = await read_excel_sheet(file_path, sheet_name)
        
        # Convert to list of dictionaries
        data = df.to_dict('records')
        columns = df.columns.tolist()
        
        return ExcelDataResponse(
            data=data,
            columns=columns,
            count=len(data)
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error reading Excel file {file_name}: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to read Excel file: {error_message}")

# Add a new endpoint for BSE alerts data
@app.get("/bse-alerts", response_model=SEBIAnalysisDataResponse)
async def get_bse_alerts_data(limit: int = 100, offset: int = 0):
    """Get BSE alerts data from the notifications database"""
    try:
        # Define path to the notifications database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "notifications.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="BSE alerts database file not found")
        
        # Connect to the database and fetch data
        def fetch_bse_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # First, get the total count of records that match our criteria
            cursor.execute("""
                SELECT COUNT(*) 
                FROM DailyLogs 
                WHERE Link IS NOT NULL AND Link != 'NIL'
            """)
            total_count = cursor.fetchone()[0]
            
            # Fetch data from DailyLogs table with limit and offset for pagination
            # Only include records where Link is not NULL and not 'NIL'
            cursor.execute("""
                SELECT SrNo, EntityName, Link, Nature, Summary, Date 
                FROM DailyLogs 
                WHERE Link IS NOT NULL AND Link != 'NIL'
                ORDER BY Date DESC, SrNo ASC 
                LIMIT ? OFFSET ?
            """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Get column names
            column_names = [description[0] for description in cursor.description]
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                # Create a dictionary with the expected keys for the frontend
                record = dict(zip(column_names, row))
                
                # Rename keys to match the frontend expectations
                record['id'] = record.pop('SrNo', None)
                record['date_key'] = record.pop('Date', '')
                record['row_index'] = record.pop('SrNo', 0)  # Use SrNo as row_index
                record['pdf_link'] = record.pop('Link', '')
                record['summary'] = record.pop('Summary', '')
                record['inserted_at'] = record.pop('Date', '')  # Use Date as inserted_at
                # Preserve EntityName and Nature for BSE alerts
                if 'EntityName' in record:
                    record['entity_name'] = record.pop('EntityName')
                else:
                    record['entity_name'] = None
                if 'Nature' in record:
                    record['nature'] = record.pop('Nature')
                else:
                    record['nature'] = None
                
                data.append(record)
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_bse_data)
        
        return SEBIAnalysisDataResponse(
            data=data,
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching BSE alerts data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch BSE alerts data: {error_message}")

# Add endpoint for SEBI analysis data
@app.get("/sebi-analysis-data", response_model=SEBIAnalysisDataResponse)
async def get_sebi_excel_data(limit: int = 100, offset: int = 0):
    """Get SEBI analysis data from the SEBI database"""
    try:
        # Define the path to the SEBI database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "sebi_excel_master.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="SEBI database file not found")
        
        # Connect to the database and fetch data
        def fetch_sebi_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # First, get the total count of records
            cursor.execute("SELECT COUNT(*) FROM excel_summaries")
            total_count = cursor.fetchone()[0]
            
            # Fetch data from excel_summaries table with limit and offset for pagination
            cursor.execute("""
                SELECT id, date_key, row_index, pdf_link, summary, inserted_at 
                FROM excel_summaries 
                ORDER BY date_key DESC, row_index ASC 
                LIMIT ? OFFSET ?
            """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                record = {
                    'id': row[0],
                    'date_key': row[1],
                    'row_index': row[2],
                    'pdf_link': row[3],
                    'summary': row[4],
                    'inserted_at': row[5]
                }
                data.append(record)
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_sebi_data)
        
        return SEBIAnalysisDataResponse(
            data=data,
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching SEBI analysis data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch SEBI analysis data: {error_message}")

# Add endpoint for RBI analysis data
@app.get("/rbi-analysis-data", response_model=SEBIAnalysisDataResponse)
async def get_rbi_excel_data(limit: int = 100, offset: int = 0):
    """Get RBI analysis data from the RBI database"""
    try:
        # Define the path to the RBI database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "rbi.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="RBI database file not found")
        
        # Connect to the database and fetch data
        def fetch_rbi_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # First, get the total count of records that match our criteria
            cursor.execute("""
                SELECT COUNT(*) 
                FROM master_summaries 
                WHERE NOT (pdf_link = 'NIL' AND summary = 'NIL')
            """)
            total_count = cursor.fetchone()[0]
            
            # Fetch data from master_summaries table with limit and offset for pagination
            # Only exclude records where both pdf_link and summary are 'NIL'
            cursor.execute("""
                SELECT id, run_date, pdf_link, summary, created_at 
                FROM master_summaries 
                WHERE NOT (pdf_link = 'NIL' AND summary = 'NIL')
                ORDER BY run_date DESC, id ASC 
                LIMIT ? OFFSET ?
            """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                record = {
                    'id': row[0],
                    'date_key': row[1],  # Using date_key to match frontend expectations
                    'row_index': row[0],  # Using id as row_index
                    'pdf_link': row[2],
                    'summary': row[3],
                    'inserted_at': row[4]
                }
                data.append(record)
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_rbi_data)
        
        return SEBIAnalysisDataResponse(
            data=data,
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching RBI analysis data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch RBI analysis data: {error_message}")

# Add admin authentication endpoints
@app.post("/admin/login", response_model=AdminLoginResponse)
async def admin_login(credentials: AdminLoginRequest):
    """Authenticate admin user"""
    try:
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and verify credentials
        def verify_admin_credentials():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Check if the provided credentials match any admin user
            cursor.execute("""
                SELECT id, username FROM admin_credentials 
                WHERE username = ? AND password = ?
            """, (credentials.username, credentials.password))
            
            result = cursor.fetchone()
            conn.close()
            
            return result
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        admin_user = await loop.run_in_executor(thread_pool, verify_admin_credentials)
        
        if admin_user:
            # In a real application, you would generate a proper JWT token
            # For now, we'll just return a success response
            return AdminLoginResponse(
                success=True,
                message="Login successful",
                token=f"admin_token_{admin_user[0]}"  # Simple token for demonstration
            )
        else:
            return AdminLoginResponse(
                success=False,
                message="Invalid credentials"
            )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error during admin login: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to authenticate: {error_message}")

# Add email management endpoints (admin only)
@app.get("/emails", response_model=EmailListResponse)
async def get_emails(search: Optional[str] = None):
    """Get all email addresses with optional search filter"""
    try:
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and fetch emails
        def fetch_emails():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Build query with optional search filter
            if search:
                # Case-insensitive search in email column
                query = "SELECT email FROM email WHERE LOWER(email) LIKE LOWER(?) ORDER BY email"
                search_pattern = f"%{search}%"
                cursor.execute(query, (search_pattern,))
            else:
                # Get all emails
                cursor.execute("SELECT email FROM email ORDER BY email")
            
            rows = cursor.fetchall()
            conn.close()
            
            return [row[0] for row in rows]
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        emails = await loop.run_in_executor(thread_pool, fetch_emails)
        
        return EmailListResponse(
            emails=emails,
            count=len(emails)
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching emails: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch emails: {error_message}")

@app.post("/emails", response_model=EmailResponse)
async def add_email(email_entry: EmailEntry):
    """Add a new email address (admin only)"""
    try:
        # Validate email format
        import re
        email_regex = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
        if not re.match(email_regex, email_entry.email):
            raise HTTPException(status_code=400, detail="Invalid email format")
        
        # Validate that email is from adani.com or pspprojects.com domain (case-insensitive)
        email_lower = email_entry.email.lower()
        if not email_lower.endswith('@adani.com') and not email_lower.endswith('@pspprojects.com'):
            raise HTTPException(status_code=400, detail="Only emails from adani.com or pspprojects.com domains are allowed")
        
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and add email
        def add_email_to_db():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Check if email already exists (case-insensitive)
            cursor.execute("SELECT COUNT(*) FROM email WHERE LOWER(email) = LOWER(?)", (email_entry.email,))
            count = cursor.fetchone()[0]
            
            if count > 0:
                conn.close()
                raise HTTPException(status_code=409, detail="Email already exists")
            
            # Insert new email
            cursor.execute("INSERT INTO email (email) VALUES (?)", (email_entry.email,))
            conn.commit()
            conn.close()
            
            return email_entry.email
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        email = await loop.run_in_executor(thread_pool, add_email_to_db)
        
        return EmailResponse(email=email)
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error adding email: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to add email: {error_message}")

@app.delete("/emails/{email_address}", response_model=EmailResponse)
async def delete_email(email_address: str):
    """Delete an email address (admin only)"""
    try:
        # Decode URL encoded email address
        email = urllib.parse.unquote(email_address)
        
        # Define path to the email database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "email_data.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Database file not found")
        
        # Connect to the database and delete email
        def delete_email_from_db():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Check if email exists
            cursor.execute("SELECT COUNT(*) FROM email WHERE email = ?", (email,))
            count = cursor.fetchone()[0]
            
            if count == 0:
                conn.close()
                raise HTTPException(status_code=404, detail="Email not found")
            
            # Delete email
            cursor.execute("DELETE FROM email WHERE email = ?", (email,))
            conn.commit()
            conn.close()
            
            return email
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        email = await loop.run_in_executor(thread_pool, delete_email_from_db)
        
        return EmailResponse(email=email)
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error deleting email: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to delete email: {error_message}")

# Add endpoint to get directors data
@app.get("/directors", response_model=DirectorsResponse)
async def get_directors(search: str = "", limit: int = 100, offset: int = 0):
    """Get directors data from the directors database"""
    try:
        # Define path to the directors database file
        db_path = os.path.join(os.path.dirname(__file__), "public", "directors.db")
        
        # Check if database file exists
        if not os.path.exists(db_path):
            raise HTTPException(status_code=404, detail="Directors database file not found")
        
        # Connect to the database and fetch data
        def fetch_directors_data():
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            
            # Build query based on search parameter
            search_term = f"%{search}%" if search else ""
            if search:
                # Search in both name and DIN fields
                query = """
                    SELECT id, name, din, created_at 
                    FROM directors 
                    WHERE name LIKE ? OR din LIKE ?
                    ORDER BY name
                    LIMIT ? OFFSET ?
                """
                cursor.execute(query, (search_term, search_term, limit, offset))
            else:
                # Get all directors
                cursor.execute("""
                    SELECT id, name, din, created_at 
                    FROM directors 
                    ORDER BY name
                    LIMIT ? OFFSET ?
                """, (limit, offset))
            
            rows = cursor.fetchall()
            
            # Convert to list of dictionaries
            data = []
            for row in rows:
                record = {
                    'id': row[0],
                    'name': row[1],
                    'din': row[2],
                    'created_at': row[3]
                }
                data.append(record)
            
            # Get total count
            if search:
                cursor.execute("SELECT COUNT(*) FROM directors WHERE name LIKE ? OR din LIKE ?", (search_term, search_term))
            else:
                cursor.execute("SELECT COUNT(*) FROM directors")
            total_count = cursor.fetchone()[0]
            
            conn.close()
            return data, total_count
        
        # Run the database operation in a thread pool
        loop = asyncio.get_event_loop()
        data, total_count = await loop.run_in_executor(thread_pool, fetch_directors_data)
        
        return DirectorsResponse(
            data=[DirectorResponse(**record) for record in data],
            count=total_count
        )
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error fetching directors data: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch directors data: {error_message}")

# Add API endpoints for place management
@app.get("/places", response_model=PlacesResponse)
async def get_places():
    """Get all places from the places database"""
    try:
        conn = get_places_db()
        cursor = conn.cursor()
        
        cursor.execute('SELECT id, name, address, is_default, created_at FROM places ORDER BY is_default DESC, name')
        rows = cursor.fetchall()
        
        places = []
        for row in rows:
            places.append({
                "id": row[0],
                "name": row[1],
                "address": row[2],
                "is_default": bool(row[3]),
                "created_at": row[4]
            })
        
        conn.close()
        return {"data": places, "count": len(places)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching places: {str(e)}")

@app.post("/places", response_model=PlaceResponse)
async def create_place(place: PlaceCreateRequest):
    """Create a new place"""
    try:
        conn = get_places_db()
        cursor = conn.cursor()
        
        # If this is set as default, unset the current default
        if place.is_default:
            cursor.execute('UPDATE places SET is_default = 0')
        
        cursor.execute('''
            INSERT INTO places (name, address, is_default)
            VALUES (?, ?, ?)
        ''', (place.name, place.address, place.is_default))
        
        place_id = cursor.lastrowid
        conn.commit()
        conn.close()
        
        return {
            "id": place_id,
            "name": place.name,
            "address": place.address,
            "is_default": place.is_default,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating place: {str(e)}")

@app.put("/places/{place_id}", response_model=PlaceResponse)
async def update_place(place_id: int, place: PlaceUpdateRequest):
    """Update an existing place"""
    try:
        conn = get_places_db()
        cursor = conn.cursor()
        
        # Check if place exists
        cursor.execute('SELECT id, name, address, is_default, created_at FROM places WHERE id = ?', (place_id,))
        row = cursor.fetchone()
        if not row:
            conn.close()
            raise HTTPException(status_code=404, detail="Place not found")
        
        # If setting as default, unset current default
        if place.is_default:
            cursor.execute('UPDATE places SET is_default = 0')
        
        # Update fields that are provided
        update_fields = []
        update_values = []
        
        if place.name is not None:
            update_fields.append("name = ?")
            update_values.append(place.name)
        
        if place.address is not None:
            update_fields.append("address = ?")
            update_values.append(place.address)
        
        if place.is_default is not None:
            update_fields.append("is_default = ?")
            update_values.append(place.is_default)
        
        if update_fields:
            update_values.append(place_id)
            cursor.execute(f'''
                UPDATE places 
                SET {", ".join(update_fields)}
                WHERE id = ?
            ''', update_values)
        
        conn.commit()
        
        # Fetch updated place
        cursor.execute('SELECT id, name, address, is_default, created_at FROM places WHERE id = ?', (place_id,))
        updated_row = cursor.fetchone()
        conn.close()
        
        if not updated_row:
            raise HTTPException(status_code=404, detail="Place not found after update")
        
        return {
            "id": updated_row[0],
            "name": updated_row[1],
            "address": updated_row[2],
            "is_default": bool(updated_row[3]),
            "created_at": updated_row[4]
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating place: {str(e)}")

@app.delete("/places/{place_id}")
async def delete_place(place_id: int):
    """Delete a place"""
    try:
        conn = get_places_db()
        cursor = conn.cursor()
        
        # Check if place exists
        cursor.execute('SELECT id FROM places WHERE id = ?', (place_id,))
        row = cursor.fetchone()
        if not row:
            conn.close()
            raise HTTPException(status_code=404, detail="Place not found")
        
        # Check if this is the default place
        cursor.execute('SELECT is_default FROM places WHERE id = ?', (place_id,))
        is_default = cursor.fetchone()[0]
        
        # If deleting the default place, we should prevent it or set another as default
        if is_default:
            # Count total places
            cursor.execute('SELECT COUNT(*) FROM places')
            total_places = cursor.fetchone()[0]
            if total_places <= 1:
                conn.close()
                raise HTTPException(status_code=400, detail="Cannot delete the last place or the default place when it's the only one")
        
        cursor.execute('DELETE FROM places WHERE id = ?', (place_id,))
        conn.commit()
        conn.close()
        
        return {"message": "Place deleted successfully"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting place: {str(e)}")

# Add endpoint to refresh directors database from Excel file
@app.post("/directors/refresh")
async def refresh_directors_db():
    """Refresh the directors database from the Excel file"""
    try:
        # Define paths
        backend_dir = os.path.dirname(os.path.abspath(__file__))
        excel_file = os.path.join(backend_dir, "public", "List of Directors.xlsx")
        db_file = os.path.join(backend_dir, "public", "directors.db")
        
        # Check if Excel file exists
        if not os.path.exists(excel_file):
            raise HTTPException(status_code=404, detail="Excel file not found")
        
        # Read the Excel file
        df = pd.read_excel(excel_file)
        
        # Connect to the SQLite database
        conn = sqlite3.connect(db_file)
        cursor = conn.cursor()
        
        # Create table if it doesn't exist
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS directors (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                email TEXT NOT NULL,
                phone TEXT NOT NULL,
                address TEXT NOT NULL,
                city TEXT NOT NULL,
                state TEXT NOT NULL,
                zip TEXT NOT NULL,
                country TEXT NOT NULL,
                website TEXT NOT NULL,
                description TEXT NOT NULL
            )
        ''')
        
        # Insert data into the table
        for index, row in df.iterrows():
            cursor.execute('''
                INSERT INTO directors (name, email, phone, address, city, state, zip, country, website, description)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (row['Name'], row['Email'], row['Phone'], row['Address'], row['City'], row['State'], row['Zip'], row['Country'], row['Website'], row['Description']))
        
        # Commit changes and close the connection
        conn.commit()
        conn.close()
        
        return {"message": "Directors database refreshed successfully"}
    except Exception as e:
        return {"error": str(e)}

# Add a helper function to get database connection for places
def get_places_db():
    """Get database connection for places"""
    backend_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)))
    db_file = os.path.join(backend_dir, "public", "places.db")
    return sqlite3.connect(db_file)
        df = pd.read_excel(excel_file, sheet_name="Sheet2")
        
        # Clean the data
        df = df.dropna(subset=['Name', 'DIN'])
        df['DIN'] = df['DIN'].astype(str)
        
        # Create database connection
        conn = sqlite3.connect(db_file)
        cursor = conn.cursor()
        
        # Insert directors data
        inserted_count = 0
        for _, row in df.iterrows():
            try:
                cursor.execute(
                    "INSERT OR REPLACE INTO directors (name, din) VALUES (?, ?)",
                    (row['Name'], str(row['DIN']))
                )
                inserted_count += 1
            except Exception as e:
                logger.error(f"Error inserting director {row['Name']}: {e}")
        
        # Commit changes and close connection
        conn.commit()
        conn.close()
        
        return {"message": f"Successfully refreshed directors database with {inserted_count} directors"}
    except HTTPException:
        raise
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error refreshing directors database: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to refresh directors database: {error_message}")

def format_date_for_document(date_str: str) -> str:
    """Format ISO date string (YYYY-MM-DD) to document format (DD Month YYYY)"""
    try:
        if isinstance(date_str, str) and date_str:
            # Parse ISO date
            date_obj = datetime.strptime(date_str, "%Y-%m-%d")
            month_names = ["January", "February", "March", "April", "May", "June",
                          "July", "August", "September", "October", "November", "December"]
            # Get ordinal day (1st, 2nd, 3rd, etc.)
            day = date_obj.day
            if 10 <= day % 100 <= 20:
                suffix = "th"
            else:
                suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return f"{day}{suffix} {month_names[date_obj.month - 1]} {date_obj.year}"
        return date_str
    except Exception as e:
        logger.warning(f"Error formatting date {date_str}: {e}")
        return date_str

def format_date_simple(date_str: str) -> str:
    """Format ISO date string to simple format (DD Month YYYY) without ordinal"""
    try:
        if isinstance(date_str, str) and date_str:
            date_obj = datetime.strptime(date_str, "%Y-%m-%d")
            month_names = ["January", "February", "March", "April", "May", "June",
                          "July", "August", "September", "October", "November", "December"]
            return f"{date_obj.day:02d} {month_names[date_obj.month - 1]} {date_obj.year}"
        return date_str
    except Exception as e:
        logger.warning(f"Error formatting date {date_str}: {e}")
        return date_str

def format_ordinal(num: int) -> str:
    """Convert number to ordinal (1 -> 1st, 2 -> 2nd, 3 -> 3rd, etc.)"""
    if 10 <= num % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(num % 10, "th")
    return f"{num}{suffix}"

def format_auditor_amount(amount: int) -> str:
    """Format auditor payment amount with Indian number formatting"""
    amount_str = str(amount)
    if len(amount_str) <= 3:
        return amount_str
    # Add commas for Indian numbering (lakhs, crores)
    result = []
    amount_str = amount_str[::-1]  # Reverse for easier processing
    for i, digit in enumerate(amount_str):
        if i == 3 or (i > 3 and (i - 3) % 2 == 0):
            result.append(',')
        result.append(digit)
    return ''.join(result[::-1])

def replace_placeholders_in_text(text: str, minutes_data: Q1MinutesData) -> str:
    """Replace all Q1 template placeholders in text with proper formatting"""
    if not text:
        return text
    
    # Format dates
    meeting_date_formatted = format_date_simple(minutes_data.meetingDate)
    previous_minutes_date_formatted = format_date_simple(minutes_data.previousMinutesDate)
    recording_date_formatted = format_date_simple(minutes_data.recordingDate)
    signing_date_formatted = format_date_simple(minutes_data.signingDate)
    
    # Format auditor amount
    auditor_amount_formatted = format_auditor_amount(minutes_data.auditorPaymentNumber)
    
    # Create director and attendee lists
    # Format directors with Chairman label for the first one if it matches chairmanName
    present_directors_list = []
    for idx, director in enumerate(minutes_data.presentDirectors):
        director_line = f"Mr. {director.name} (DIN: {director.din})"
        # If this is the chairman and it's the first director, mark as Chairman
        if idx == 0 and director.name == minutes_data.chairmanName:
            director_line += "\tChairman"
        elif idx > 0 or director.name != minutes_data.chairmanName:
            director_line += "\tDirector"
        present_directors_list.append(director_line)
    present_directors_str = "\n".join(present_directors_list) if present_directors_list else ""
    
    interest_disclosures_list = [f"Mr. {director.name} (DIN: {director.din})" for director in minutes_data.interestDisclosures]
    interest_disclosures_str = "\n".join(interest_disclosures_list) if interest_disclosures_list else ""
    
    disqualification_declarations_list = [f"Mr. {director.name} (DIN: {director.din})" for director in minutes_data.disqualificationDeclarations]
    disqualification_declarations_str = "\n".join(disqualification_declarations_list) if disqualification_declarations_list else ""
    
    in_attendance_list = [f"{attendee.name} ({attendee.role})" for attendee in minutes_data.inAttendance]
    in_attendance_str = "\n".join(in_attendance_list) if in_attendance_list else ""
    
    # Replacement mapping - order matters! Replace longer/more specific patterns first
    replacements = {
        # Company & Meeting Header - Specific patterns first
        "[Time: COMMENCED AT]": minutes_data.timeCommenced,
        "[Time: CONCLUDED AT]": minutes_data.timeConcluded,
        "[Date of Meeting]": meeting_date_formatted,
        "[Day of Meeting]": minutes_data.meetingDay,
        "[Place of Meeting]": minutes_data.meetingPlace,
        "[Name of Company]": minutes_data.companyName,
        "[No. of Meeting]": minutes_data.meetingNumber,
        "[Type of Meeting]": minutes_data.meetingType,
        
        # Attendance & Directors - Handle [from MCA] carefully
        # Note: We'll handle [from MCA] contextually, not here in the main replacement
        "[From website: MCA]": minutes_data.registeredOfficeAddress,
        "[From Website: MCA]": minutes_data.registeredOfficeAddress,  # Case variation
        "[FROM WEBSITE: MCA]": minutes_data.registeredOfficeAddress,  # Case variation
        
        # Chairman and Manual entries (context-specific) - Order matters!
        # MUST replace specific patterns BEFORE generic [Manual]
        "Mr. [Manual] occupied the Chair": f"Mr. {minutes_data.chairmanName} occupied the Chair",
        "In ATTENDANCE: – [Manual]": f"In ATTENDANCE: – {in_attendance_str}" if in_attendance_str else "In ATTENDANCE: – [Authorised Officer to be added]",
        "In ATTENDANCE: - [Manual]": f"In ATTENDANCE: - {in_attendance_str}" if in_attendance_str else "In ATTENDANCE: - [Authorised Officer to be added]",
        "In ATTENDANCE – [Manual]": f"In ATTENDANCE – {in_attendance_str}" if in_attendance_str else "In ATTENDANCE – [Authorised Officer to be added]",
        
        # Previous minutes
        "[Auto]": previous_minutes_date_formatted,
        
        # Auditor Payment - Specific patterns first (BEFORE generic [Manual] replacement)
        "Rs. [Manual]/-": f"Rs. {auditor_amount_formatted}/-",
        "Rs [Manual]/-": f"Rs {auditor_amount_formatted}/-",  # Without period
        "(Rupees [Manual] Only)": f"(Rupees {minutes_data.auditorPaymentWords} Only)",
        "Rupees [Manual] Only": f"Rupees {minutes_data.auditorPaymentWords} Only",  # Without parentheses
        "31ST MARCH, 20____": f"31ST MARCH, {minutes_data.auditorPaymentYear}",
        "31st March, 20____": f"31st March, {minutes_data.auditorPaymentYear}",
        "31st March 20____": f"31st March {minutes_data.auditorPaymentYear}",
        "31ST MARCH 20____": f"31ST MARCH {minutes_data.auditorPaymentYear}",
        "31 March 20____": f"31 March {minutes_data.auditorPaymentYear}",  # Without ordinal
        "31 MARCH 20____": f"31 MARCH {minutes_data.auditorPaymentYear}",
        # Handle 3 underscores variant (common in templates)
        "31ST MARCH, 20___": f"31ST MARCH, {minutes_data.auditorPaymentYear}",
        "31st March, 20___": f"31st March, {minutes_data.auditorPaymentYear}",
        "31st March 20___": f"31st March {minutes_data.auditorPaymentYear}",
        # Handle any remaining 20___ patterns
        "20___": str(minutes_data.auditorPaymentYear),
        
        # Financial Statements - Year placeholders
        "20<year>": f"20{minutes_data.directorsReportYear}",
        "<year>": str(minutes_data.fsYear),
        # Directors' Report year - ensure full year is used (handle various patterns)
        "31ST MARCH, 2020": f"31ST MARCH, {minutes_data.directorsReportYear}",  # Fix hardcoded 2020
        "31st March, 2020": f"31st March, {minutes_data.directorsReportYear}",
        "31st march, 2020": f"31st march, {minutes_data.directorsReportYear}",
        "31ST MARCH, 2020 ": f"31ST MARCH, {minutes_data.directorsReportYear}",  # With trailing space
        "31st March, 2020 ": f"31st March, {minutes_data.directorsReportYear}",
        
        # Related Party Transactions - Handle various formats
        "<2025 to  2026>__": f"{minutes_data.rptFinYearRange.from_year} to {minutes_data.rptFinYearRange.to_year}",
        "<2024 to  2025>__": f"{minutes_data.rptFinYearRange.from_year} to {minutes_data.rptFinYearRange.to_year}",
        "<2025 to 2026>__": f"{minutes_data.rptFinYearRange.from_year} to {minutes_data.rptFinYearRange.to_year}",  # Single space
        "<2024 to 2025>__": f"{minutes_data.rptFinYearRange.from_year} to {minutes_data.rptFinYearRange.to_year}",
        
        # Signatory patterns - Very specific patterns first (handle variations)
        "Mr. <Director Name>, __, Director (DIN: _<DIN Number>___)": f"Mr. {minutes_data.signatory1.name}, {minutes_data.signatory1.role} (DIN: {minutes_data.signatory1.din})",
        "Mr. <Director Name>, __, Director (DIN: <DIN Number>___)": f"Mr. {minutes_data.signatory1.name}, {minutes_data.signatory1.role} (DIN: {minutes_data.signatory1.din})",
        "Mr. __<Director Name>__, Director (DIN: __<DIN>__)": f"Mr. {minutes_data.signatory2.name}, Director (DIN: {minutes_data.signatory2.din})",
        "Mr. <Director Name>__, Director (DIN: __<DIN>__)": f"Mr. {minutes_data.signatory2.name}, Director (DIN: {minutes_data.signatory2.din})",
        
        # AGM Details - Specific patterns with variations (with ordinal formatting)
        "the ____ Day of ____": f"the {format_ordinal(minutes_data.agmDay)} Day of {minutes_data.agmMonth}",
        "the ____ day of ____": f"the {format_ordinal(minutes_data.agmDay)} day of {minutes_data.agmMonth}",
        "____ Day of ____": f"{format_ordinal(minutes_data.agmDay)} Day of {minutes_data.agmMonth}",
        "20____ at ____ p.m.": f"{minutes_data.agmYear} at {minutes_data.agmTime}",
        "20____ at ____ P.M.": f"{minutes_data.agmYear} at {minutes_data.agmTime}",
        "20____ at ____ pm": f"{minutes_data.agmYear} at {minutes_data.agmTime}",
        "20____ at ____ PM": f"{minutes_data.agmYear} at {minutes_data.agmTime}",
        "<no of meeting>": minutes_data.agmNumber,
        "<no. of meeting>": minutes_data.agmNumber,  # With period
        "<day>": minutes_data.agmDayName,
        "<Chairman>": minutes_data.chairmanShortName,
        "Mr. <Chairman>": f"Mr. {minutes_data.chairmanShortName}",  # With prefix
        "Mr. <Chairman>, Chairman": f"Mr. {minutes_data.chairmanShortName}, Chairman",
        # Handle hardcoded "Mr. Mehta" in template
        "Mr. Mehta": f"Mr. {minutes_data.chairmanShortName.replace('Mr. ', '').replace('Mr ', '')}" if minutes_data.chairmanShortName.startswith('Mr.') or minutes_data.chairmanShortName.startswith('Mr ') else f"Mr. {minutes_data.chairmanShortName}",
        "Mr. Mehta,": f"Mr. {minutes_data.chairmanShortName.replace('Mr. ', '').replace('Mr ', '')}," if minutes_data.chairmanShortName.startswith('Mr.') or minutes_data.chairmanShortName.startswith('Mr ') else f"Mr. {minutes_data.chairmanShortName},",
        "Mehta": minutes_data.chairmanShortName.replace('Mr. ', '').replace('Mr ', ''),
        
        # Signing Information - Specific patterns with variations
        "Date of Recording    :\t____": f"Date of Recording    :\t{recording_date_formatted}",
        "Date of Recording : ____": f"Date of Recording : {recording_date_formatted}",
        "Date of Recording: ____": f"Date of Recording: {recording_date_formatted}",  # Without space
        "Date of Recording    : ____": f"Date of Recording    : {recording_date_formatted}",  # With spaces instead of tab
        "Date of Signing\t:\t____": f"Date of Signing\t:\t{signing_date_formatted}",
        "Date of Signing : ____": f"Date of Signing : {signing_date_formatted}",
        "Date of Signing: ____": f"Date of Signing: {signing_date_formatted}",  # Without space
        "Date of Signing\t:": f"Date of Signing\t:\t{signing_date_formatted}",  # Just colon, no placeholder
        "Date of Signing :": f"Date of Signing : {signing_date_formatted}",  # Just colon with space
        "Date of Signing:": f"Date of Signing: {signing_date_formatted}",  # Just colon, no space
        "Place\t\t\t: \t____\t\t\t\tChairman": f"Place\t\t\t: \t{minutes_data.signingPlace}\t\t\t\tChairman",
        "Place : ____": f"Place : {minutes_data.signingPlace}",
        "Place: ____": f"Place: {minutes_data.signingPlace}",  # Without space
    }
    
    # Apply replacements in order
    for placeholder, replacement in replacements.items():
        if placeholder in text:
            text = text.replace(placeholder, replacement)
    
    # Replace generic [Manual] placeholder AFTER all specific patterns (only if not already replaced)
    # This handles cases where [Manual] appears in other contexts
    if "[Manual]" in text:
        text = text.replace("[Manual]", minutes_data.chairmanName)
    
    # Clean up duplicate DIN patterns in disclosure sections BEFORE replacing [from MCA]
    # This pattern appears when [from MCA] gets replaced multiple times
    duplicate_pattern = re.compile(
        r'Mr\.\s+Mr\.\s+([^(]+)\s+\(DIN:\s+([^)]+)\)\s+\(DIN:\s+Mr\.\s+\1\s+\(DIN:\s+\2\)\)',
        re.IGNORECASE
    )
    text = duplicate_pattern.sub(r'Mr. \1 (DIN: \2)', text)
    
    # Handle interest disclosures section (Section 184) - Use regex for flexible matching
    # This should happen AFTER we've replaced [from MCA] in the main replacement
    # But we need to handle cases where the pattern still exists
    if interest_disclosures_str:
        # Pattern for Section 184 disclosures - handle various formats
        patterns = [
            r'Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)',
            r'Mr\.\s*\[from\s*MCA\]\s*\(DIN:\s*\[from\s*MCA\]\)',  # With spaces
        ]
        for pattern in patterns:
            text = re.sub(pattern, interest_disclosures_str, text, flags=re.IGNORECASE)
    
    # Handle disqualification declarations section (Section 164(2))
    if disqualification_declarations_str:
        # Similar pattern for Section 164(2)
        patterns = [
            r'Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)',
            r'Mr\.\s*\[from\s*MCA\]\s*\(DIN:\s*\[from\s*MCA\]\)',
        ]
        for pattern in patterns:
            text = re.sub(pattern, disqualification_declarations_str, text, flags=re.IGNORECASE)
    
    # Clean up any remaining year placeholders with context
    # Replace 20____ with appropriate year (try fsYear first, then others)
    text = re.sub(r'20____', str(minutes_data.fsYear), text)
    # Also replace 20___ (3 underscores) - common variant
    text = re.sub(r'20___', str(minutes_data.fsYear), text)
    
    # Fix hardcoded "2020" year in Directors' Report to use correct year
    # Pattern: "31st March, 2020" -> "31st March, {directorsReportYear}"
    text = re.sub(r'31st\s+March,\s+2020\s*', f'31st March, {minutes_data.directorsReportYear} ', text, flags=re.IGNORECASE)
    text = re.sub(r'31ST\s+MARCH,\s+2020\s*', f'31ST MARCH, {minutes_data.directorsReportYear} ', text, flags=re.IGNORECASE)
    
    # Fix malformed year patterns like "202025 _" -> "2025"
    # This handles cases where the year got concatenated incorrectly
    text = re.sub(r'20(\d{2})(\d{2,})\s*_?\s*', r'20\1 ', text)  # Fix "202025 _" -> "2025 "
    text = re.sub(r'(\d{4})(\d{2,})\s*_?\s*', lambda m: f"{m.group(1)[:4]} ", text)  # Fix any 6+ digit years
    
    # Fix year patterns with trailing underscores or spaces
    text = re.sub(r'(\d{4})\s*_+\s*', r'\1 ', text)  # Fix "2025___" -> "2025 "
    text = re.sub(r'(\d{4})\s+_+\s*', r'\1 ', text)  # Fix "2025 ___" -> "2025 "
    
    # Only remove standalone ____ if they're not part of a larger pattern
    # Be careful - don't remove underscores that are part of formatting
    text = re.sub(r'\b____\b', '', text)  # Only if it's a word boundary
    
    # Clean up any double spaces that might have been created
    text = re.sub(r'  +', ' ', text)
    
    # Fix duplicate "BOARD OF DIRECTORS MEETING" in header
    # This happens when meetingType already includes "Board of Directors Meeting"
    # Pattern: "Meeting BOARD OF DIRECTORS MEETING BOARD OF DIRECTORS MEETING"
    text = re.sub(r'([A-Z\s]+)\s+BOARD OF DIRECTORS MEETING\s+BOARD OF DIRECTORS MEETING', 
                  r'\1 BOARD OF DIRECTORS MEETING', text, flags=re.IGNORECASE)
    # Also handle case where meeting type duplicates it
    text = re.sub(r'(Board of Directors Meeting|BOARD OF DIRECTORS MEETING)\s+(Board of Directors Meeting|BOARD OF DIRECTORS MEETING)', 
                  r'\1', text, flags=re.IGNORECASE)
    
    # Fix signing date if it's just a colon with nothing after (empty placeholder)
    # This handles cases where the placeholder was removed but date wasn't filled
    if "Date of Signing" in text and signing_date_formatted:
        # Look for patterns like "Date of Signing\t:" or "Date of Signing :" with nothing after
        text = re.sub(r'Date of Signing[\s\t]*:[\s\t]*$', 
                     f'Date of Signing\t:\t{signing_date_formatted}', 
                     text, flags=re.MULTILINE)
        text = re.sub(r'Date of Signing[\s\t]*:[\s\t]*(?=\n|$)', 
                     f'Date of Signing\t:\t{signing_date_formatted}', 
                     text, flags=re.MULTILINE)
    
    return text

# Add endpoint to generate minutes document for Q1 template
@app.post("/generate-minutes/q1")
async def generate_minutes_q1(minutes_data: Q1MinutesData):
    """Generate a minutes document from Q1 template and data with standardized placeholder handling"""
    try:
        # Define paths - try multiple possible template filenames
        templates_dir = os.path.join(os.path.dirname(__file__), "public", "templates")
        possible_template_names = [
            "q1_meeting_template.docx",
            "Q1_meeting_template.docx",
            "Q1 Meeting Template.docx",
            "q1 template.docx",
            "Q1 Template.docx"
        ]
        
        template_path = None
        for template_name in possible_template_names:
            candidate_path = os.path.join(templates_dir, template_name)
            if os.path.exists(candidate_path):
                template_path = candidate_path
                break
        
        if not template_path:
            # List available templates for debugging
            available_templates = []
            if os.path.exists(templates_dir):
                available_templates = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]
            
            raise HTTPException(
                status_code=404, 
                detail=f"Q1 template file not found. Tried: {possible_template_names}. Available templates: {available_templates}"
            )
        
        # Load the template document
        doc = Document(template_path)
        logger.info(f"Loaded Q1 template from: {template_path}")
        
        # Generate output filename
        safe_meeting_number = minutes_data.meetingNumber.replace(' ', '_').replace('/', '_')
        safe_date = minutes_data.meetingDate.replace('-', '')
        filename = f"meeting_minutes_Q1_{safe_date}_{safe_meeting_number}.docx"
        output_path = os.path.join(templates_dir, filename)
        
        # Track context for disclosure sections and collect all placeholders for logging
        in_interest_section = False
        in_disqualification_section = False
        found_placeholders = set()
        missing_placeholders = set()
        
        # Helper function to replace placeholders in paragraph (handles runs)
        def replace_in_paragraph(paragraph):
            """Replace placeholders in paragraph, handling runs properly"""
            if not paragraph.text:
                return
            
            # Get full text first for context detection
            full_text = paragraph.text
            text_lower = full_text.lower()
            
            # Update context
            nonlocal in_interest_section, in_disqualification_section
            if "section 184" in text_lower or ("interest" in text_lower and "disclosure" in text_lower):
                in_interest_section = True
                in_disqualification_section = False
            elif "section 164" in text_lower or "disqualification" in text_lower or "164(2)" in text_lower:
                in_disqualification_section = True
                in_interest_section = False
            elif "attendance" in text_lower or "directors were present" in text_lower:
                in_interest_section = False
                in_disqualification_section = False
            
            # Collect placeholders for logging
            placeholder_patterns = [
                r'\[[^\]]+\]',  # [placeholder]
                r'<[^>]+>',     # <placeholder>
                r'20_{3,4}',    # Year placeholder (3 or 4 underscores)
                r'_{4}',        # Blank placeholder (4 underscores)
            ]
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, full_text)
                found_placeholders.update(matches)
            
            # Get director strings for use in context-aware replacement
            present_directors_list_ctx = [f"Mr. {director.name} (DIN: {director.din})" for director in minutes_data.presentDirectors]
            
            # Sequential replacement: first "Chairman", then "Director"
            for idx, director in enumerate(minutes_data.presentDirectors):
                if idx == 0 and director.name == minutes_data.chairmanName:
                    present_directors_list_ctx[0] += "\tChairman"
                else:
                    present_directors_list_ctx[idx] += "\tDirector"
            present_directors_str_ctx = "\n".join(present_directors_list_ctx) if present_directors_list_ctx else ""
            
            # Replace placeholders with context awareness
            # Handle attendance section first - format directors properly
            is_attendance_paragraph = (
                "directors were present" in text_lower or 
                "present:" in text_lower or
                (text_lower.strip().startswith("the following") and "directors" in text_lower)
            )
            
            if is_attendance_paragraph:
                # In attendance section, replace [from MCA] with formatted directors list
                # The template has [from MCA] on separate lines - one for Chairman, one for Director
                if "[from MCA]" in full_text and minutes_data.presentDirectors:
                    # Count how many [from MCA] placeholders we have
                    placeholder_count = full_text.count("[from MCA]")
                    
                    # First, replace [from MCA] followed by Chairman
                    chairman_found = False
                    for idx, director in enumerate(minutes_data.presentDirectors):
                        if idx == 0 and director.name == minutes_data.chairmanName and not chairman_found:
                            if re.search(r'\[from MCA\]\s+Chairman', full_text, re.IGNORECASE):
                                director_line = f"Mr. {director.name} (DIN: {director.din})\tChairman"
                                full_text = re.sub(r'\[from MCA\]\s+Chairman', director_line, full_text, count=1, flags=re.IGNORECASE)
                                chairman_found = True
                                break
                    
                    # Then replace remaining [from MCA] with Director
                    for idx, director in enumerate(minutes_data.presentDirectors):
                        if idx == 0 and director.name == minutes_data.chairmanName and chairman_found:
                            continue  # Already handled
                        director_line = f"Mr. {director.name} (DIN: {director.din})\tDirector"
                        full_text = re.sub(r'\[from MCA\]\s+Director', director_line, full_text, count=1, flags=re.IGNORECASE)
                    
                    # Replace any remaining [from MCA] with the full list (fallback)
                    if "[from MCA]" in full_text:
                        full_text = full_text.replace("[from MCA]", present_directors_str_ctx)
                
                # Do normal replacement for other placeholders
                full_text = replace_placeholders_in_text(full_text, minutes_data)
                
            elif in_interest_section and minutes_data.interestDisclosures:
                # Create numbered list format for disclosures
                disclosure_lines = []
                for idx, d in enumerate(minutes_data.interestDisclosures, 1):
                    disclosure_lines.append(f"{idx}.\tMr. {d.name} (DIN: {d.din})")
                interest_disclosures_str = "\n\n".join(disclosure_lines)
                
                # Handle whitespace variations (tabs, spaces)
                # Replace each occurrence of the pattern with corresponding list item
                # Handle duplicate numbering patterns first (1. 1. -> 1.)
                duplicate_number_pattern = re.compile(r'(\d+)\.\s+(\d+)\.\s*Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)', re.IGNORECASE)
                def replace_with_item_dup(match):
                    num = int(match.group(2))  # Use the second number
                    if 1 <= num <= len(minutes_data.interestDisclosures):
                        return f"{num}.\tMr. {minutes_data.interestDisclosures[num-1].name} (DIN: {minutes_data.interestDisclosures[num-1].din})"
                    return match.group(0)
                full_text = duplicate_number_pattern.sub(replace_with_item_dup, full_text)
                
                # Handles both space and tab separators
                # Handle patterns with tab separator: "1.\t1.\t"
                duplicate_number_pattern_tab = re.compile(r'(\d+)\.\s*\t\s*(\d+)\.\s*\t\s*Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)', re.IGNORECASE)
                def replace_with_item_dup_tab(match):
                    num = int(match.group(2))  # Use the second number
                    if 1 <= num <= len(minutes_data.interestDisclosures):
                        return f"{num}.\tMr. {minutes_data.interestDisclosures[num-1].name} (DIN: {minutes_data.interestDisclosures[num-1].din})"
                    return match.group(0)
                full_text = duplicate_number_pattern_tab.sub(replace_with_item_dup_tab, full_text)
                
                # Then handle any remaining unnumbered patterns - replace with full list
                unnumbered_pattern = re.compile(r'Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)', re.IGNORECASE)
                if unnumbered_pattern.search(full_text):
                    # Replace first occurrence with full list
                    full_text = unnumbered_pattern.sub(interest_disclosures_str, full_text, count=1)
                    # Remove any remaining occurrences
                    full_text = unnumbered_pattern.sub('', full_text)
                
                # Now do normal replacement for other placeholders
                full_text = replace_placeholders_in_text(full_text, minutes_data)
                
            elif in_disqualification_section and minutes_data.disqualificationDeclarations:
                # Create numbered list format for declarations
                declaration_lines = []
                for idx, d in enumerate(minutes_data.disqualificationDeclarations, 1):
                    declaration_lines.append(f"{idx}.\tMr. {d.name} (DIN: {d.din})")
                disqualification_str = "\n\n".join(declaration_lines)
                
                # Replace each occurrence of the pattern with corresponding list item
                # Pattern: "1.\t1.\tMr. [from MCA] (DIN: [from MCA])" or "1.\tMr. [from MCA] (DIN: [from MCA])"
                # First, handle duplicate numbering patterns (1.\t1. -> 1.)
                duplicate_number_pattern = re.compile(r'(\d+)\.\s+(\d+)\.\s*Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)', re.IGNORECASE)
                def replace_with_item_dup(match):
                    num = int(match.group(2))  # Use the second number
                    if 1 <= num <= len(minutes_data.disqualificationDeclarations):
                        return f"{num}.\tMr. {minutes_data.disqualificationDeclarations[num-1].name} (DIN: {minutes_data.disqualificationDeclarations[num-1].din})"
                    return match.group(0)
                full_text = duplicate_number_pattern.sub(replace_with_item_dup, full_text)
                
                # Handles both space and tab separators
                # Handle patterns with tab separator: "1.\t1.\t"
                duplicate_number_pattern_tab = re.compile(r'(\d+)\.\s*\t\s*(\d+)\.\s*\t\s*Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)', re.IGNORECASE)
                def replace_with_item_dup_tab(match):
                    num = int(match.group(2))  # Use the second number
                    if 1 <= num <= len(minutes_data.disqualificationDeclarations):
                        return f"{num}.\tMr. {minutes_data.disqualificationDeclarations[num-1].name} (DIN: {minutes_data.disqualificationDeclarations[num-1].din})"
                    return match.group(0)
                full_text = duplicate_number_pattern_tab.sub(replace_with_item_dup_tab, full_text)
                
                # Then handle any remaining unnumbered patterns - replace with full list
                unnumbered_pattern = re.compile(r'Mr\.\s*\[from MCA\]\s*\(DIN:\s*\[from MCA\]\)', re.IGNORECASE)
                if unnumbered_pattern.search(full_text):
                    # Replace first occurrence with full list
                    full_text = unnumbered_pattern.sub(disqualification_str, full_text, count=1)
                    # Remove any remaining occurrences
                    full_text = unnumbered_pattern.sub('', full_text)
                
                # Now do normal replacement for other placeholders
                full_text = replace_placeholders_in_text(full_text, minutes_data)
                
            else:
                # Normal replacement - handle [from MCA] if not already handled
                # Fallback for placeholders without labels
                if "[from MCA]" in full_text and present_directors_str_ctx:
                    # Default: use present directors for [from MCA]
                    full_text = full_text.replace("[from MCA]", present_directors_str_ctx)
                full_text = replace_placeholders_in_text(full_text, minutes_data)
            
            # Clean up any duplicate patterns that might have been created
            # Pattern: "Mr. Mr. Name (DIN: x) (DIN: Mr. Name (DIN: x))"
            duplicate_din_pattern = re.compile(
                r'Mr\.\s+Mr\.\s+([^(]+)\s+\(DIN:\s+([^)]+)\)\s+\(DIN:\s+Mr\.\s+\1\s+\(DIN:\s+\2\)\)',
                re.IGNORECASE
            )
            full_text = duplicate_din_pattern.sub(r'Mr. \1 (DIN: \2)', full_text)
            
            # Disclosure sections duplicate numbering fix
            # Pattern: "1. 1. Mr. Name (DIN: x)" -> "1. Mr. Name (DIN: x)"
            duplicate_numbering_pattern = re.compile(r'(\d+)\.\s+(\d+)\.\s+(Mr\.\s+[^(]+\s+\(DIN:\s+[^)]+\))')
            full_text = duplicate_numbering_pattern.sub(r'\1. \3', full_text)
            
            # Added regex patterns for tab-separated duplicates (e.g., "1.\t1.\t")
            duplicate_numbering_tab_pattern = re.compile(r'(\d+)\.\s*\t\s*(\d+)\.\s*\t\s+(Mr\.\s+[^(]+\s+\(DIN:\s+[^)]+\))')
            full_text = duplicate_numbering_tab_pattern.sub(r'\1.\t\3', full_text)
            
            # Year placeholders - Added patterns for 3 underscores (20___) and 4 underscores (20____)
            full_text = re.sub(r'20_{3,4}', str(minutes_data.fsYear), full_text)
            
            # Regex fixes for hardcoded "2020" in Directors' Report
            # Pattern: "2020" -> correct year
            full_text = re.sub(r'2020(?=\s+to\s+202[0-9])', str(minutes_data.directorsReportYear), full_text)
            full_text = re.sub(r'2020(?=\s+to\s+20[0-9]{2})', str(minutes_data.directorsReportYear), full_text)
            
            # Mr. Mehta replacement - Added patterns for "Mr. Mehta", "Mr. Mehta,", and "Mehta"
            # Only replace if the chairman name is not "Mehta" to avoid self-replacement
            if minutes_data.chairmanName.lower() != "mehta":
                mehta_patterns = [
                    r'Mr\.\s*Mehta(?=\s+occupied)',
                    r'Mr\.\s*Mehta,',
                    r'\bMehta\b',
                ]
                for pattern in mehta_patterns:
                    full_text = re.sub(pattern, minutes_data.chairmanName, full_text, flags=re.IGNORECASE)
            
            # Check for remaining placeholders
            for pattern in placeholder_patterns:
                matches = re.findall(pattern, full_text)
                missing_placeholders.update(matches)
            
            # Update paragraph text (this handles runs automatically)
            paragraph.text = full_text
        
        # Process paragraphs in the document
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        
        # Process headers and footers
        for section in doc.sections:
            # Process header
            if section.header:
                for paragraph in section.header.paragraphs:
                    replace_in_paragraph(paragraph)
            
            # Process footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    replace_in_paragraph(paragraph)
        
        # Process tables if they exist
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        
        # Log any remaining placeholders
        if missing_placeholders:
            logger.warning(f"Found {len(missing_placeholders)} placeholders that may not have been replaced: {sorted(missing_placeholders)}")
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Generated Q1 minutes document: {filename}")
        
        # Prepare response
        response = {
            "success": True,
            "filename": filename,
            "message": "Minutes document generated successfully",
            "template_used": os.path.basename(template_path)
        }
        
        # Add placeholder information if any are missing
        if missing_placeholders:
            response["warning"] = f"Some placeholders may not have been replaced: {sorted(missing_placeholders)}"
            response["missing_placeholders"] = sorted(missing_placeholders)
        
        if found_placeholders:
            response["found_placeholders"] = sorted(found_placeholders)
        
        return response
    except HTTPException:
        raise
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error generating Q1 minutes document: {error_message}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate minutes document: {error_message}")

# Add endpoint to generate minutes document for other templates
@app.post("/generate-minutes/other")
async def generate_minutes_other(minutes_data: MinutesData):
    """Generate a minutes document from other templates and data"""
    try:
        # Define paths - use lowercase for template filename
        template_filename = f"{minutes_data.template.lower()}_meeting_template.docx"
        template_path = os.path.join(os.path.dirname(__file__), "public", "templates", template_filename)
        output_dir = os.path.join(os.path.dirname(__file__), "public", "templates")
        
        # Check if template exists
        if not os.path.exists(template_path):
            # List available templates for debugging
            templates_dir = os.path.join(os.path.dirname(__file__), "public", "templates")
            available_templates = []
            if os.path.exists(templates_dir):
                available_templates = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]
            
            raise HTTPException(
                status_code=404, 
                detail=f"Template file {template_filename} not found. Available templates: {available_templates}"
            )
        
        # Load the template document
        doc = Document(template_path)
        
        # Generate output filename
        filename = f"meeting_minutes_{minutes_data.template}_{minutes_data.meetingDate.replace('-', '')}_{minutes_data.meetingNumber.replace(' ', '_')}.docx"
        output_path = os.path.join(output_dir, filename)
        
        # Process other templates with existing field structure
        # Create a string of director names with DINs
        directors_str = "\n".join([f"Mr. {director.name} (DIN: {director.din})" for director in minutes_data.directors])
        
        # Create a string of present directors (if provided)
        present_directors_str = "\n".join([f"Mr. {director.name} (DIN: {director.din})" for director in minutes_data.presentDirectors])
        if not present_directors_str:
            present_directors_str = directors_str  # Fallback to directors list
        
        # Create a string of attendees
        in_attendance_str = "\n".join([f"{attendee.name} ({attendee.role})" for attendee in minutes_data.inAttendance])
        if not in_attendance_str:
            in_attendance_str = minutes_data.authorisedOfficer  # Fallback to authorised officer
        
        # Process the document content
        for paragraph in doc.paragraphs:
            # Replace placeholders with actual values
            text = paragraph.text
            text = text.replace("[Name of Company]", minutes_data.companyName)
            text = text.replace("[No. of Meeting]", minutes_data.meetingNumber)
            text = text.replace("[Type of Meeting]", minutes_data.meetingType)
            text = text.replace("[Day of Meeting]", minutes_data.meetingDay)
            text = text.replace("[Date of Meeting]", minutes_data.meetingDate)
            text = text.replace("[Time: COMMENCED AT]", minutes_data.meetingStartTime)
            text = text.replace("[Time: CONCLUDED AT]", minutes_data.meetingEndTime)
            text = text.replace("[Place of Meeting]", minutes_data.meetingPlace)
            text = text.replace("[Manual]", minutes_data.chairmanName)
            text = text.replace("[Auto]", minutes_data.previousMeetingDate)
            text = text.replace("20____", str(minutes_data.financialYear))
            text = text.replace("[from MCA]", present_directors_str)  # Use present directors or fallback
            text = text.replace("[Authorised Officer]", in_attendance_str)  # Use in attendance or fallback
            text = text.replace("[Quorum details]", minutes_data.quorum)
            text = text.replace("[Concern/Interest details]", minutes_data.concerns)
            text = text.replace("[Declaration details]", minutes_data.declarations)
            text = text.replace("[AGM Number]", minutes_data.agmNumber)
            text = text.replace("[AGM Day]", str(minutes_data.agmDay))
            text = text.replace("[AGM Date]", minutes_data.agmDate)
            text = text.replace("[AGM Time]", minutes_data.agmTime)
            text = text.replace("[AGM Year]", str(minutes_data.financialYear))
            text = text.replace("[From website: MCA]", minutes_data.agmPlace)
            text = text.replace("[Recording Date]", minutes_data.recordingDate)
            text = text.replace("[Signing Date]", minutes_data.signingDate)
            text = text.replace("[Signing Place]", minutes_data.signingPlace)  # Add signing place
            
            # Handle exact signing information placeholders from Q2 template
            if minutes_data.recordingDate:
                text = text.replace("Date of Recording    :\t____", f"Date of Recording    :\t{minutes_data.recordingDate}")
            # For Date of Signing, the template doesn't have ____ but expects the date after the colon
            if minutes_data.signingDate:
                text = text.replace("Date of Signing\t:\t", f"Date of Signing\t:\t{minutes_data.signingDate}")
            if minutes_data.signingPlace:
                text = text.replace("Place\t\t\t: \t____\t\t\t\tChairman", f"Place\t\t\t: \t{minutes_data.signingPlace}\t\t\t\tChairman")
            
            # Update the paragraph text
            paragraph.text = text
        
        # Also process tables if they exist
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        text = text.replace("[Name of Company]", minutes_data.companyName)
                        text = text.replace("[No. of Meeting]", minutes_data.meetingNumber)
                        text = text.replace("[Type of Meeting]", minutes_data.meetingType)
                        text = text.replace("[Day of Meeting]", minutes_data.meetingDay)
                        text = text.replace("[Date of Meeting]", minutes_data.meetingDate)
                        text = text.replace("[Time: COMMENCED AT]", minutes_data.meetingStartTime)
                        text = text.replace("[Time: CONCLUDED AT]", minutes_data.meetingEndTime)
                        text = text.replace("[Place of Meeting]", minutes_data.meetingPlace)
                        text = text.replace("[Manual]", minutes_data.chairmanName)
                        text = text.replace("[Auto]", minutes_data.previousMeetingDate)
                        text = text.replace("20____", str(minutes_data.financialYear))
                        text = text.replace("[from MCA]", present_directors_str)  # Use present directors or fallback
                        text = text.replace("[Authorised Officer]", in_attendance_str)  # Use in attendance or fallback
                        text = text.replace("[Quorum details]", minutes_data.quorum)
                        text = text.replace("[Concern/Interest details]", minutes_data.concerns)
                        text = text.replace("[Declaration details]", minutes_data.declarations)
                        text = text.replace("[AGM Number]", minutes_data.agmNumber)
                        text = text.replace("[AGM Day]", str(minutes_data.agmDay))
                        text = text.replace("[AGM Date]", minutes_data.agmDate)
                        text = text.replace("[AGM Time]", minutes_data.agmTime)
                        text = text.replace("[AGM Year]", str(minutes_data.financialYear))
                        text = text.replace("[From website: MCA]", minutes_data.agmPlace)
                        text = text.replace("[Recording Date]", minutes_data.recordingDate)
                        text = text.replace("[Signing Date]", minutes_data.signingDate)
                        text = text.replace("[Signing Place]", minutes_data.signingPlace)  # Add signing place
                        
                        # Handle exact signing information placeholders from Q2 template
                        if minutes_data.recordingDate:
                            text = text.replace("Date of Recording    :\t____", f"Date of Recording    :\t{minutes_data.recordingDate}")
                        # For Date of Signing, the template doesn't have ____ but expects the date after the colon
                        if minutes_data.signingDate:
                            text = text.replace("Date of Signing\t:\t", f"Date of Signing\t:\t{minutes_data.signingDate}")
                        if minutes_data.signingPlace:
                            text = text.replace("Place\t\t\t: \t____\t\t\t\tChairman", f"Place\t\t\t: \t{minutes_data.signingPlace}\t\t\t\tChairman")
                        
                        # Update the paragraph text
                        paragraph.text = text
        
        # Save the document
        doc.save(output_path)
        
        # Return the path to the generated file
        return {
            "success": True,
            "filename": filename,
            "message": "Minutes document generated successfully"
        }
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error generating minutes document: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to generate minutes document: {error_message}")

# Keep the original endpoint for backward compatibility but fix the logic
@app.post("/generate-minutes")
async def generate_minutes(request: Request):
    """Generate a minutes document from template and data"""
    try:
        # Get raw JSON data
        raw_data = await request.json()
        
        # Check which template type we're using
        template = raw_data.get("template", "Q1")
        
        if template == "Q1":
            # Validate with Q1MinutesData model
            minutes_data = Q1MinutesData(**raw_data)
            return await generate_minutes_q1(minutes_data)
        else:
            # Validate with existing MinutesData model for other templates
            minutes_data = MinutesData(**raw_data)
            return await generate_minutes_other(minutes_data)
            
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error generating minutes document: {error_message}")
        raise HTTPException(status_code=500, detail=f"Failed to generate minutes document: {error_message}")

# Thread pool for CPU-intensive operations
thread_pool = concurrent.futures.ThreadPoolExecutor(max_workers=4)

# Custom static files handler for SPA
class SPAStaticFiles(StaticFiles):
    async def get_response(self, path: str, scope):
        try:
            response = await super().get_response(path, scope)
import os
import sqlite3
from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.exceptions import HTTPException

app = FastAPI()

# Add a helper function to get database connection for places
def get_places_db():
    """Get database connection for places"""
    backend_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)))
    db_file = os.path.join(backend_dir, "public", "places.db")
    return sqlite3.connect(db_file)

# Serve static files from the dist directory (React build) - this must be the last route
# Only add this if the dist directory exists
DIST_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "dist")
if os.path.exists(DIST_DIR):
    app.mount("/", SPAStaticFiles(directory=DIST_DIR, html=True), name="static")

# Add a test endpoint to verify static file serving is working
@app.get("/test-static")
async def test_static_serving():
    """Test endpoint to verify static file serving is working"""
    return {"message": "Static file serving should be working correctly"}

if __name__ == "__main__":
    import uvicorn
    
    # Run without SSL on port 8001
    uvicorn.run(app, host="0.0.0.0", port=8000)
